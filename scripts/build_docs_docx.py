#!/usr/bin/env python3
"""Build Word (.docx) from Markdown guides in docs/. Requires python-docx."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError as exc:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from exc


def strip_backticks(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def add_runs_with_bold(paragraph, text: str) -> None:
    """Split by **bold** markers and add runs; `code` shown without backticks."""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        m = re.match(r"\*\*(.+?)\*\*$", part)
        if m:
            run = paragraph.add_run(strip_backticks(m.group(1)))
            run.bold = True
        else:
            paragraph.add_run(strip_backticks(part))


_IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")


def _add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell_text = row[c_idx].strip() if c_idx < len(row) else ""
            cell_text = strip_backticks(re.sub(r"^\*+|\*+$", "", cell_text))
            table.rows[r_idx].cells[c_idx].text = cell_text


def markdown_to_docx(md_path: Path, docx_path: Path, doc_title: str) -> Path:
    doc = Document()
    cp = doc.core_properties
    cp.title = doc_title
    cp.language = "ru-RU"

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    md_dir = md_path.parent
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        i += 1

        if not raw.strip():
            continue
        if raw.strip() == "---":
            continue

        img_match = _IMAGE_RE.match(raw.strip())
        if img_match:
            caption, rel_path = img_match.group(1).strip(), img_match.group(2).strip()
            image_path = (md_dir / rel_path).resolve()
            if image_path.is_file():
                try:
                    doc.add_picture(str(image_path), width=Inches(6.2))
                except Exception as exc:
                    doc.add_paragraph(f"[Изображение не вставлено: {image_path.name} — {exc}]")
            else:
                doc.add_paragraph(f"[Изображение не найдено: {rel_path}]")
            if caption:
                cap = doc.add_paragraph()
                cap_run = cap.add_run(caption)
                cap_run.italic = True
                cap_run.font.size = Pt(9)
            continue

        if raw.strip().startswith("|") and raw.strip().endswith("|"):
            table_rows: list[list[str]] = []
            while i <= len(lines):
                row_line = raw.strip()
                if not (row_line.startswith("|") and row_line.endswith("|")):
                    break
                if re.match(r"^\|[\s\-:|]+\|$", row_line):
                    raw = lines[i].rstrip() if i < len(lines) else ""
                    i += 1
                    continue
                cells = [cell.strip() for cell in row_line.strip("|").split("|")]
                table_rows.append(cells)
                if i >= len(lines):
                    break
                raw = lines[i].rstrip()
                i += 1
            _add_markdown_table(doc, table_rows)
            if raw.strip() and not (raw.strip().startswith("|") and raw.strip().endswith("|")):
                i -= 1
            continue

        if raw.startswith("# "):
            doc.add_heading(strip_backticks(raw[2:].strip()), level=0)
            continue
        if raw.startswith("## "):
            doc.add_heading(strip_backticks(raw[3:].strip()), level=1)
            continue
        if raw.startswith("### "):
            doc.add_heading(strip_backticks(raw[4:].strip()), level=2)
            continue
        if raw.startswith("#### "):
            doc.add_heading(strip_backticks(raw[5:].strip()), level=3)
            continue

        if raw.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, raw.strip()[2:].strip())
            continue

        num = re.match(r"^(\s*)(\d+)\.\s+(.*)$", raw)
        if num:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_bold(p, num.group(3).strip())
            continue

        p = doc.add_paragraph()
        add_runs_with_bold(p, raw.strip())

    try:
        doc.save(str(docx_path))
    except PermissionError:
        alt = docx_path.with_name(docx_path.stem + ".generated" + docx_path.suffix)
        doc.save(str(alt))
        print(f"Note: {docx_path} is locked; wrote {alt} instead. Close Word and re-run.", file=sys.stderr)
        return alt
    return docx_path


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    docs = root / "docs"
    pairs = [
        (docs / "INTERFACE_GUIDE_USER_RU.md", docs / "INTERFACE_GUIDE_USER_RU.docx", "CrewDeck — инструкция пользователя"),
        (docs / "INTERFACE_GUIDE_RU.md", docs / "INTERFACE_GUIDE_RU.docx", "CrewDeck — полная документация интерфейса"),
        (
            docs / "PODACHA_INSTRUCTION_STAFF_RU.md",
            docs / "PODACHA_INSTRUCTION_STAFF_RU.docx",
            "CrewDeck — ПОДАЧА (инструкция для сотрудников)",
        ),
    ]
    full_manual = root / "docs" / "reports" / "USER_MANUAL_FULL_WORKFLOW_RU.docx"
    gen_script = root / "scripts" / "generate_user_manual_full_docx.py"
    if gen_script.is_file():
        import subprocess

        subprocess.run([sys.executable, str(gen_script)], check=False, cwd=str(root))
        if full_manual.is_file():
            print(f"Wrote {full_manual}")
    for md, out, title in pairs:
        if not md.is_file():
            print(f"Skip missing {md}", file=sys.stderr)
            continue
        path = markdown_to_docx(md, out, title)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
