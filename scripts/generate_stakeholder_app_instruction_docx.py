"""Generate the Manisov stakeholder application instruction DOCX.

Run with the bundled workspace Python:
python scripts/generate_stakeholder_app_instruction_docx.py
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "docs" / "reports" / "MANISOV_STAKEHOLDER_APP_INSTRUCTION_RU.docx"

ACCENT = "2E74B5"
ACCENT_DARK = "1F4D78"
INK = "1F2933"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
WARN_FILL = "FFF4D6"
OK_FILL = "EAF7EF"
BORDER = "D0D7DE"


SCREENSHOTS = [
    (
        "Dashboard: главное меню CRM",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "01-dashboard.png",
        "Стартовый экран после входа: быстрые переходы к Seamens Data, Notifications и другим разделам.",
    ),
    (
        "Seamens Data: список кандидатов",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "02-candidates-list.png",
        "Основной рабочий список: поиск, фильтры, пагинация, загрузка анкет и переход в карточку кандидата.",
    ),
    (
        "Карточка кандидата: панель действий",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "03-toolbar-podacha.png",
        "Панель верхних действий карточки: ПОДАЧА, генерация документов, контракт и административные действия.",
    ),
    (
        "ПОДАЧА: окно сборки пакета",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "04-podacha-modal.png",
        "Окно выбора opening/previous vessel, шаблонов DOCX и сканов для ZIP-пакета.",
    ),
    (
        "ПОДАЧА: выбор судов",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "05-podacha-vessels.png",
        "Поля opening m/v и previous m/v используются только для текущего ZIP и не меняют карточку кандидата.",
    ),
    (
        "ПОДАЧА: шаблоны",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "06-podacha-templates.png",
        "CRM берет шаблоны из Templates, подставляет данные кандидата и кладет готовые DOCX в архив.",
    ),
    (
        "ПОДАЧА: сканы документов",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "07-podacha-scans-documents.png",
        "В архив можно добавить только те документы, где уже есть прикрепленный скан.",
    ),
    (
        "ПОДАЧА: сборка ZIP",
        PROJECT_ROOT / "docs" / "screenshots" / "podacha" / "08-podacha-build-zip.png",
        "После нажатия кнопки система формирует ZIP с info-list и выбранными сканами.",
    ),
    (
        "Contract Tab: макет вкладки контракта",
        PROJECT_ROOT / "docs" / "reports" / "contract_tab_mockup.png",
        "Макет показывает целевой UX для контрактных данных и синхронизации с salary calculation.",
    ),
    (
        "Contract Tab: генерация контрактного документа",
        PROJECT_ROOT / "docs" / "reports" / "contract_generate_modal_mockup.png",
        "Макет окна генерации контрактного DOCX из выбранного шаблона.",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_dxa = int(sum(widths_inches) * 1440)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    if old_grid is not None:
        table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for index, width in enumerate(widths_inches):
            cell = row.cells[index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_paragraph_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, ACCENT_DARK, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Manisov CRM - stakeholder instruction")
    set_run_font(run, 8.5, MUTED)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("Manisov CRM")
    set_run_font(run, 26, ACCENT_DARK, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("Инструкция для стейкхолдера по приложению")
    set_run_font(run, 16, INK, True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    run = meta.add_run(f"Версия: {date.today().isoformat()} | Среда: production + GitHub main | Формат: stakeholder handover")
    set_run_font(run, 10.5, MUTED)

    add_callout(
        doc,
        "Назначение документа",
        "Документ дает стейкхолдеру полную картину продукта: что уже реализовано, как пользоваться основными сценариями, "
        "как устроены данные и деплой, что нужно проверять при приемке и какие ограничения остаются. Пароли, токены и секреты не включены.",
        fill=LIGHT_FILL,
    )


def add_paragraph(doc: Document, text: str, *, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        set_run_font(r, bold=True)
        p.add_run(" ")
    run = p.add_run(text)
    set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        run = p.add_run(item)
        set_run_font(run)


def add_callout(doc: Document, label: str, body: str, *, fill: str = LIGHT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=BORDER, size="4")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label)
    set_run_font(run, 11, ACCENT_DARK, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    set_run_font(run2, 10.5, INK)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, 10.5, ACCENT_DARK, True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, 10, INK)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("24292F")
    doc.add_paragraph()


def add_image(doc: Document, title: str, path: Path, caption: str) -> None:
    heading = doc.add_paragraph(style="Heading 3")
    set_paragraph_keep_with_next(heading)
    heading.add_run(title)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        try:
            run = p.add_run()
            run.add_picture(str(path), width=Inches(6.15))
        except Exception as exc:
            warn = p.add_run(f"[Скриншот не вставлен: {path.name}; {exc}]")
            set_run_font(warn, 10, "9B1C1C")
    else:
        p = doc.add_paragraph()
        warn = p.add_run(f"[Файл скриншота не найден: {path}]")
        set_run_font(warn, 10, "9B1C1C")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    set_run_font(run, 9.5, MUTED)


def add_section(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_section(doc, "1. Краткое резюме", 1)
    add_paragraph(
        doc,
        "Manisov CRM - производственная CRM для ведения базы моряков, их анкет, документов, сертификатов, "
        "морского стажа, компаний, судов, шаблонов Word и пакетов ПОДАЧА. Приложение построено на FastAPI, React/Vite, "
        "PostgreSQL и Docker Compose. Production сервер самостоятельно подтягивает изменения из GitHub main через systemd timer.",
    )
    add_table(
        doc,
        ["Область", "Текущее состояние для стейкхолдера"],
        [
            ["Продукт", "CRM уже развернута на production и доступна через frontend/API порты сервера."],
            ["Данные", "Используется PostgreSQL 16 в Docker volume; файлы хранятся в uploads, generated и templates."],
            ["Деплой", "Основной механизм - server-side systemd timer, который проверяет GitHub main и пересобирает сервисы при новом коммите."],
            ["Качество", "Backend автотесты проходили: 286 passed, 1 skipped; frontend build выполнялся успешно."],
            ["Ограничения", "GitHub Actions не является активным production-деплоем без secrets; HTTPS/domain/Nginx и backup policy стоит оформить отдельным этапом."],
        ],
        [1.45, 5.05],
    )
    add_callout(
        doc,
        "Важно для принятия решения",
        "Система уже пригодна для операционной приемки: можно проверять вход, роли, список кандидатов, карточку, документы, ПОДАЧА, шаблоны, "
        "компании/суда и уведомления. Для промышленной эксплуатации с внешними пользователями нужно отдельно закрыть HTTPS, домен, backup и security hardening.",
        fill=OK_FILL,
    )

    add_section(doc, "2. Доступы, окружения и ответственность", 1)
    add_table(
        doc,
        ["Пункт", "Описание"],
        [
            ["GitHub", "Основной репозиторий: bernik1986/Manisov, рабочая ветка production flow: main."],
            ["Production", "Сервер: 178.212.12.161, путь приложения: /var/www/Manisov."],
            ["Frontend", "Ожидаемый endpoint production: http://178.212.12.161:5173."],
            ["Backend API", "Ожидаемый endpoint production: http://178.212.12.161:8000/docs."],
            ["Секреты", "Пароли и JWT секреты хранятся на сервере в .env и не должны попадать в Git или документы."],
            ["Admin", "Production admin credentials хранятся на сервере отдельно; в этот документ они не включаются."],
        ],
        [1.45, 5.05],
    )
    add_paragraph(
        doc,
        "Владельцу продукта важно понимать: любые изменения в коде должны проходить через GitHub main. Production не редактируется вручную как источник правды; "
        "сервер только подтягивает и запускает то, что лежит в репозитории.",
    )

    add_section(doc, "3. Карта приложения", 1)
    add_table(
        doc,
        ["Раздел", "Для кого", "Что делает"],
        [
            ["Login", "Все пользователи", "Вход по логину/паролю, выдача JWT токена, защита закрытых разделов."],
            ["Dashboard / Menu", "Все пользователи", "Главное меню, сводка по кандидатам и уведомлениям, быстрые переходы."],
            ["Seamens Data", "Recruiter, Admin, Viewer", "Список кандидатов, загрузка анкет/CV, поиск, фильтры по должности и флоту, открытие карточки."],
            ["Candidate Card", "Recruiter, Admin, Viewer", "Полная карточка кандидата: profile, recruitment, documents, certificates, sea service, scans, family contacts."],
            ["Templates", "Recruiter, Admin", "Файловый менеджер DOC/DOCX/PDF шаблонов для генерации документов и ПОДАЧА."],
            ["Company", "Recruiter, Admin", "Справочник компаний, судов, плейсхолдеров и salary templates."],
            ["Notifications", "Все роли", "Истекающие документы, отсутствие сканов, переход в нужную строку карточки."],
            ["Users", "Admin", "Создание пользователей, смена ролей/паролей, блокировка и удаление учеток."],
            ["Logs", "Admin", "Audit log: кто и когда менял данные."],
        ],
        [1.4, 1.45, 3.65],
    )

    add_section(doc, "4. Роли и права доступа", 1)
    add_table(
        doc,
        ["Роль", "Доступ"],
        [
            ["admin", "Полный доступ: кандидаты, загрузки, редактирование, удаление, пользователи, audit logs, справочники."],
            ["recruiter", "Операционная работа: кандидаты, документы, сканы, ПОДАЧА, шаблоны, компании и суда без управления пользователями."],
            ["viewer", "Просмотр данных без редактирования, без Users/Logs и без опасных действий."],
        ],
        [1.25, 5.25],
    )
    add_bullets(
        doc,
        [
            "Удаление кандидата и управление пользователями должны оставаться только у admin.",
            "Viewer полезен для менеджеров/стейкхолдеров, которым нужен контроль без риска изменить данные.",
            "Нельзя оставлять систему без активного admin: код защищает от блокировки последнего администратора.",
        ],
    )

    add_section(doc, "5. Основные бизнес-сценарии", 1)
    add_section(doc, "5.1. Новый кандидат из анкеты", 2)
    add_numbers(
        doc,
        [
            "Пользователь входит в CRM и открывает Seamens Data.",
            "Загружает анкету DOC/DOCX/PDF/XLS/XLSX или CV PDF.",
            "Backend парсит файл, нормализует поля и создает кандидата либо предлагает обновить найденный дубль.",
            "Recruiter проверяет карточку: профиль, recruitment, документы, сертификаты, sea service.",
            "Прикрепляются сканы к соответствующим строкам Documents/Certificates/Diplomas/Medical/Flag documents.",
            "Для отправки клиенту собирается ПОДАЧА ZIP или генерируются отдельные DOCX по шаблонам.",
        ],
    )

    add_section(doc, "5.2. Обновление существующего кандидата", 2)
    add_numbers(
        doc,
        [
            "Загрузить новую анкету или открыть карточку вручную.",
            "Если система определила дубликат, подтвердить обновление существующей записи.",
            "Проверить ключевые поля, потому что часть данных приходит из разных форматов анкет.",
            "Сохранить измененные группы в карточке и проверить Notifications.",
        ],
    )

    add_section(doc, "5.3. Подготовка пакета ПОДАЧА", 2)
    add_numbers(
        doc,
        [
            "Открыть карточку кандидата и нажать ПОДАЧА.",
            "Выбрать opening m/v и previous m/v при необходимости.",
            "Отметить шаблоны info-list и нужные сканы.",
            "Нажать Собрать ZIP и отправить архив клиенту вне CRM.",
        ],
    )
    add_callout(
        doc,
        "Отличие ПОДАЧА от генерации документов",
        "ПОДАЧА создает ZIP с документами и сканами. Кнопка Сгенерировать документы скачивает отдельные DOCX по выбранным шаблонам без архива и без сканов.",
        fill=LIGHT_FILL,
    )

    add_section(doc, "5.4. Управление шаблонами и справочниками", 2)
    add_bullets(
        doc,
        [
            "Templates: структура папок, загрузка DOC/DOCX/PDF, переименование, удаление, скачивание.",
            "Company: папки, компании, суда, IMO, flag, vessel type, плейсхолдеры для шаблонов.",
            "Salary templates: компоненты зарплаты по company/rank и синхронизация с кандидатской salary calculation.",
            "Contract: сохранение контрактных данных кандидата и подготовка к генерации контрактных DOCX.",
        ],
    )

    add_section(doc, "6. Скриншоты приложения", 1)
    add_paragraph(
        doc,
        "Ниже - набор актуальных экранов и макетов, которые показывают основные рабочие зоны CRM. Они нужны стейкхолдеру для приемки UX и объяснения процесса команде.",
    )
    for index, (title, path, caption) in enumerate(SCREENSHOTS, start=1):
        if index in (4, 9):
            doc.add_page_break()
        add_image(doc, f"{index}. {title}", path, caption)

    add_section(doc, "7. Данные и база данных", 1)
    add_table(
        doc,
        ["Слой", "Текущее решение"],
        [
            ["DB engine", "PostgreSQL 16-alpine в Docker Compose, контейнер maritime-db."],
            ["Persistence", "Docker volume postgres_data для базы; uploads/generated/templates монтируются как папки проекта."],
            ["ORM/schema", "SQLAlchemy models в models/schema.py, миграции Alembic в migrations/versions."],
            ["Core entities", "User, Role, Candidate, Application, Document, Certificate, FlagDocument, SeaService, FamilyContact, Attachment."],
            ["Directories", "uploads - сканы, generated - сгенерированные файлы, templates - исходные шаблоны."],
            ["Reset", "Полный reset БД допустим только вручную и только при осознанной потере данных."],
        ],
        [1.5, 5.0],
    )
    add_section(doc, "7.1. Основные сущности", 2)
    add_table(
        doc,
        ["Entity", "Назначение"],
        [
            ["Candidate", "Главная карточка моряка: персональные данные, контакты, summary fields, contract/salary JSON."],
            ["Application", "Заявка кандидата: applied rank/position, salary, proposed vessel, dates."],
            ["Document", "Паспорта, seaman books, визы и другие документы."],
            ["Certificate", "Diplomas, medical certificates, STCW certificates и прочие сертификационные строки."],
            ["SeaService", "Морской стаж: vessel, type, rank, dates, duration, reason of discharge."],
            ["Attachment", "Прикрепленные сканы к документам, сертификатам, флаг-документам и другим объектам."],
            ["Company/Vessel", "Справочник компаний и судов для плейсхолдеров и контрактной логики."],
            ["TemplateFolder/TemplateFile", "Дерево шаблонов, используемых в генерации DOCX и ПОДАЧА."],
            ["AuditLog/Notification", "Журнал действий и уведомления по срокам/сканам."],
        ],
        [1.7, 4.8],
    )

    add_section(doc, "8. Production deployment и автодеплой", 1)
    add_paragraph(
        doc,
        "Production работает через Docker Compose на сервере. В отличие от GitHub Actions, который пока не является активным production pipeline без настроенных secrets, "
        "фактический автодеплой настроен на самом сервере через systemd timer.",
    )
    add_table(
        doc,
        ["Компонент", "Описание"],
        [
            ["maritime-db", "PostgreSQL container, healthcheck через pg_isready."],
            ["maritime-backend", "FastAPI container: alembic upgrade head, затем uvicorn app.main:app на 8000."],
            ["maritime-frontend", "Vite/React container на 5173, проксирование API на backend:8000."],
            ["manisov-deploy.timer", "Периодически запускает deploy service, который сравнивает GitHub main и локальный checkout."],
            ["manisov-deploy.service", "One-shot deploy script: pull/fetch, сборка/перезапуск Docker Compose при изменениях."],
        ],
        [1.7, 4.8],
    )
    add_section(doc, "8.1. Нормальный путь изменения", 2)
    add_numbers(
        doc,
        [
            "Изменение делается локально в репозитории.",
            "Запускаются нужные тесты и build.",
            "Создается commit в main и push в GitHub.",
            "Production timer на сервере подтягивает новый commit.",
            "Docker Compose пересобирает/перезапускает измененные сервисы.",
            "Проверяются HTTP 200 frontend/backend и статус контейнеров.",
        ],
    )
    add_section(doc, "8.2. Команды проверки для IT", 2)
    add_code_block(
        doc,
        [
            "ssh root@178.212.12.161",
            "cd /var/www/Manisov",
            "git rev-parse --short HEAD",
            "docker compose ps",
            "curl -I http://127.0.0.1:5173",
            "curl -I http://127.0.0.1:8000/docs",
            "systemctl status manisov-deploy.timer --no-pager",
            "journalctl -u manisov-deploy.service -n 50 --no-pager",
        ],
    )

    add_section(doc, "9. Качество, приемка и ограничения", 1)
    add_table(
        doc,
        ["Проверка", "Статус/как читать"],
        [
            ["Backend tests", "Последняя полная проверка: 286 passed, 1 skipped."],
            ["Frontend build", "npm run build выполнялся успешно после установки зависимостей."],
            ["Timer deploy", "Проверялся фактический подтяг commit с GitHub на production; frontend/backend отдавали HTTP 200."],
            ["Security", ".env вынесен из Git, secrets обязательные, backend production command без --reload."],
            ["npm audit", "Есть предупреждения npm audit; их нужно отдельно оценить и закрыть."],
            ["GitHub Actions", "Workflow есть как потенциальный путь, но без secrets не считается активным production деплоем."],
        ],
        [1.65, 4.85],
    )
    add_section(doc, "9.1. Acceptance checklist для стейкхолдера", 2)
    add_bullets(
        doc,
        [
            "Открывается frontend production и форма login.",
            "Admin/recruiter/viewer видят только разрешенные разделы.",
            "Seamens Data показывает список кандидатов, поиск и фильтры работают.",
            "Загрузка анкеты создает или обновляет карточку.",
            "В карточке сохраняются profile/recruitment/documents/certificates/sea service/family contacts.",
            "Сканы прикрепляются, скачиваются и попадают в ПОДАЧА.",
            "ПОДАЧА формирует ZIP с info-list и выбранными сканами.",
            "Templates и Company позволяют вести справочники без правки кода.",
            "Notifications ведут на нужного кандидата/документ.",
            "После push нового commit production подтягивает изменение без ручного копирования файлов.",
        ],
    )
    add_section(doc, "9.2. Что стоит запланировать следующим этапом", 2)
    add_table(
        doc,
        ["Приоритет", "Работа"],
        [
            ["High", "Настроить HTTPS, домен и reverse proxy вместо прямого доступа к портам 5173/8000."],
            ["High", "Описать и включить регулярные backup PostgreSQL + uploads/templates."],
            ["High", "Закрыть npm audit findings или зафиксировать accepted risk."],
            ["Medium", "Выбрать один production deploy path: оставить systemd timer или настроить GitHub Actions secrets/webhook."],
            ["Medium", "Завершить чистку старых названий Parcer/CrewDeck/Maritime в UI/доках, где они еще встречаются."],
            ["Medium", "Добавить больше актуальных screenshot coverage для Users, Logs, Notifications, Templates, Company."],
        ],
        [1.2, 5.3],
    )

    add_section(doc, "10. Приложение: быстрый словарь", 1)
    add_table(
        doc,
        ["Термин", "Значение"],
        [
            ["ПОДАЧА", "ZIP-пакет для отправки кандидата клиенту: generated info-list + selected scans."],
            ["Template", "DOC/DOCX/PDF файл в Templates, из которого CRM может генерировать документы."],
            ["Placeholder", "Переменная в DOCX вида {{ surname }} или {{ company_slug_vessel_slug_imo }}."],
            ["Scan slot code", "Код документа/сертификата, который используется в именах PDF сканов."],
            ["Salary calculation", "Расчет зарплатных компонентов по company/rank/template с сохранением в карточке."],
            ["Contract JSON", "Структурные контрактные данные кандидата, сохраненные в БД для генерации контрактов."],
            ["Alembic", "Механизм версионирования схемы PostgreSQL."],
            ["systemd timer", "Linux-механизм периодического запуска deploy script на сервере."],
        ],
        [1.5, 5.0],
    )

    add_callout(
        doc,
        "Финальное правило эксплуатации",
        "Источник правды - GitHub main и данные в production PostgreSQL. Код меняем через Git, данные пользователей защищаем backup и доступами, секреты не копируем в документы и чаты.",
        fill=WARN_FILL,
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    try:
        doc.save(str(OUTPUT))
        print(f"Wrote {OUTPUT}")
    except PermissionError:
        alternate = OUTPUT.with_name(OUTPUT.stem + ".generated" + OUTPUT.suffix)
        doc.save(str(alternate))
        print(f"Wrote {alternate} because {OUTPUT.name} is locked")


if __name__ == "__main__":
    main()
