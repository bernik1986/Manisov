"""Generate stakeholder report DOCX. Run: python scripts/generate_stakeholder_report_docx.py"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "docs" / "reports" / "STAKEHOLDER_REPORT_10_DAYS_RU.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = val
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    title = doc.add_heading("Отчёт для стейкхолдеров: разработка CRM Parcer", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Период: 17–27 {date.today().strftime('%B %Y')} (10 дней)\n")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub.add_run("Продукт: CRM для подбора моряков (Seamens Data, карточка кандидата, документы, шаблоны)").font.size = Pt(11)

    doc.add_paragraph()

    add_heading(doc, "1. Краткое резюме", 1)
    add_para(
        doc,
        "За 10 дней команда доработала CRM в четырёх направлениях: (1) ПОДАЧА и работа со сканами; "
        "(2) фильтры Seamens Data по должности и флоту; (3) автоматизация в карточке кандидата "
        "(морской стаж и сертификаты); (4) тесты, документация и тестовые данные для приёмки.",
    )
    add_para(
        doc,
        "Часть изменений уже в репозитории (коммит от 21.05.2026); значительный объём реализован "
        "в рабочей ветке и готов к коммиту и выкладке на стенд.",
    )

    add_heading(doc, "2. Что сделано по блокам", 1)

    add_heading(doc, "2.1. ПОДАЧА (пакеты документов) — в репозитории", 2)
    add_table(
        doc,
        ["Функция", "Зачем бизнесу"],
        [
            ["Формирование пакетов ПОДАЧА", "Быстрая отправка комплекта документов по кандидату"],
            ["Конвертация сканов в PDF", "Единый формат для архива и пересылки"],
            ["Правила именования файлов", "Понятные имена (должность, фамилия, тип документа)"],
            ["Расширенное тестовое покрытие", "Меньше регрессий при релизах"],
        ],
    )

    add_heading(doc, "2.2. Seamens Data — фильтры «Должность» и «Флот»", 2)
    add_para(doc, "Проблема: разные формулировки в анкетах не находились фильтром или давали ложные совпадения.", bold=True)
    add_bullets(
        doc,
        [
            "Должность в списке — только из первой заявки (position_applied_for / rank_applied_for).",
            "Флот в списке — только из последнего контракта в морском стаже (vessel_type по дате посадки).",
            "Нормализация: словари синонимов, 25 канонических должностей, 16 типов судов.",
            "Исправлены баги: «CO» не цеплял Second Officer; «CE» — слово officer; фильтр совпадает с колонкой.",
        ],
    )
    add_para(doc, "Справочник флота (16 типов):", bold=True)
    add_para(
        doc,
        "Bulk Carrier, General Cargo Vessel, Container Vessel, LNG Carrier, LPG Carrier, "
        "Oil/Chemical Tanker, Chemical Tanker, Crude Oil Tanker, VLCC, Tug, Passenger Vessel, "
        "Offshore Vessel, Heavy-Lift Vessel, Reefer, Ro-Ro, Multi-Purpose Vessel — с полным списком синонимов.",
    )
    add_para(doc, "Документация: docs/SEAMENS_DATA_FILTERS_RU.md")

    add_heading(doc, "2.3. Карточка кандидата — Sea service (морской стаж)", 2)
    add_table(
        doc,
        ["Функция", "Поведение"],
        [
            ["Duration", "Автоматически из дат From/To: лет/месяцев/дней (напр. 0/4/11), оба дня включительно"],
            ["Reason of Discharge", "По умолчанию EOC, поле редактируемое"],
            ["Rank и Должность", "Одно поле БД rank_on_vessel; в UI две колонки (дублирование отображения)"],
        ],
    )

    add_heading(doc, "2.4. Карточка кандидата — Certificates (сертификаты)", 2)
    add_table(
        doc,
        ["Режим", "Описание"],
        [
            ["+5 лет (по умолчанию)", "Дата окончания = выдача + 5 календарных лет (или выдача = окончание − 5 лет)"],
            ["Unlimited", "Без даты окончания, в таблице «Unlimited»"],
            ["Другое", "Обе даты вручную; кнопка «Применить +5 лет» пересчитывает вторую дату"],
        ],
    )
    add_para(
        doc,
        "Блок доступен при добавлении и при редактировании любого сертификата (не только при неполных датах из анкеты).",
    )

    add_heading(doc, "2.5. Тестовые данные и окружение", 2)
    add_table(
        doc,
        ["Инструмент", "Назначение"],
        [
            ["seed_full_demo_candidates.py", "50 полных анкет DemoSeaman001–050 со всеми секциями и PDF-сканами"],
            ["seed_filter_test_candidates.py", "Лёгкие записи для проверки фильтров"],
            ["Docker", "Локальный стенд UI + API + БД для приёмки"],
        ],
    )

    add_heading(doc, "2.6. Документация и качество", 2)
    add_bullets(
        doc,
        [
            "Обновлены README, INTERFACE_GUIDE, USER_SCENARIOS, QA_REQUIREMENTS.",
            "70+ автотестов: фильтры, нормализация, стаж, сертификаты, миграция.",
            "Runbook: docs/normalize_migration_runbook.md",
        ],
    )

    add_heading(doc, "3. Статус поставки", 1)
    add_table(
        doc,
        ["Статус", "Содержание"],
        [
            ["В git (main)", "ПОДАЧА, сканы PDF, тесты по пакетам"],
            ["Готово локально", "Фильтры Seamens Data, 16 флотов, стаж, сертификаты, сиды, доки, тесты"],
            ["Рекомендация", "Релиз-коммит + smoke: фильтры, карточка, сертификаты, стаж"],
        ],
    )

    add_heading(doc, "4. Ценность для бизнеса", 1)
    add_para(
        doc,
        "Рекрутеры получают предсказуемый поиск по должности и типу судна, меньше ручного ввода "
        "в стаже и сертификатах, единые правила сроков (+5 лет / Unlimited) и готовый демо-набор "
        "для обучения и приёмки. Снижается риск ошибок в фильтрах и расхождений между колонкой "
        "в списке и фактической выборкой кандидатов.",
    )

    add_heading(doc, "5. Сценарий демо (5 минут)", 1)
    add_bullets(
        doc,
        [
            "Seamens Data → фильтры Fleet (16 типов) и Position → список совпадает с колонками.",
            "Карточка → Sea service → даты → автоматический Duration, EOC в причине увольнения.",
            "Certificates → добавить/редактировать → +5 лет / Unlimited / Другое.",
            "Поиск DemoSeaman → 50 тестовых анкет с документами и сканами.",
        ],
    )

    add_heading(doc, "6. Ключевые модули (для IT)", 1)
    add_bullets(
        doc,
        [
            "app/rank_normalization.py, app/fleet_normalization.py",
            "app/sea_service_duration.py, app/certificate_validity.py",
            "app/frontend: CertificateValidityControls.jsx, SeaServiceSection.jsx, CandidateList.jsx",
            "scripts: normalize_ranks_and_fleet.py, seed_full_demo_candidates.py",
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"Сформировано автоматически · Parcer CRM · {date.today().isoformat()}")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
