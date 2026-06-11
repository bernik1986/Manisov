"""Generate a Word catalog of all project DOCX placeholders.

Run with the bundled workspace Python:
python scripts/generate_placeholders_catalog_docx.py
"""

from __future__ import annotations

import ast
import re
from collections import defaultdict
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "docs" / "reports" / "MANISOV_PLACEHOLDERS_CATALOG_RU.docx"

ACCENT = "2E74B5"
ACCENT_DARK = "1F4D78"
INK = "1F2933"
MUTED = "667085"
TABLE_FILL = "E8EEF5"
BORDER = "D0D7DE"

SCHEMA_PATH = PROJECT_ROOT / "models" / "schema.py"

MODEL_CLASSES = {
    "Application": ("Цикл applications", "application"),
    "Document": ("Цикл documents", "document"),
    "Certificate": ("Цикл certificates", "certificate"),
    "FlagDocument": ("Цикл flag_documents", "flag_document"),
    "SeaService": ("Цикл sea_service", "sea_service"),
    "FamilyContact": ("Цикл family_contacts", "family_contact"),
    "Attachment": ("Цикл attachments", "attachment"),
}

CANDIDATE_EXCLUDED_FIELDS = {
    "ukr_contract_json",
    "salary_calculation_json",
    "contract_json",
}

SUFFIX_LABELS = {
    "document_number": "Номер документа",
    "certificate_number": "Номер сертификата",
    "issue_date": "Дата выдачи",
    "expiry_date": "Дата окончания",
    "issuing_authority": "Кем выдан",
    "place_of_issue": "Место выдачи",
    "country_of_issue": "Страна выдачи",
    "visa_code": "Код визы",
    "visa_name": "Название визы",
    "competency_rank": "COC Rank / звание",
    "imo": "IMO",
    "flag": "Флаг",
    "port_of_registry": "Port of Registry",
    "type": "Тип судна",
    "registry_address": "Адрес регистрации судна",
    "official_number": "Official No",
    "call_sign": "Call Sign",
    "grt": "GRT",
    "deadweight": "Dead Weight",
    "year_built": "Год постройки",
    "engine_type": "Engine Type",
    "engine_hp": "H.P.",
    "classification_society": "Classification society",
}

SALARY_LABELS = {
    "salary_company": "Компания для salary calculation",
    "salary_rank": "Должность для salary calculation",
    "salary_total_wage": "Total Wage",
    "salary_period_of_employment": "Period of employment",
    "salary_basic_monthly_wage": "Basic monthly wage",
    "salary_monthly_overtime": "Monthly overtime",
    "salary_overtime_rate": "Overtime rate",
    "salary_sepf": "SEPF",
    "salary_imtf": "IMTF",
    "salary_leave": "Leave",
    "salary_leave_sub": "Leave sub",
    "salary_various_extra_overtime": "Various extra overtime",
    "salary_fixed_components_total": "Fixed components total",
    "salary_owners_bonus": "Owners bonus",
}

CONTRACT_BASE_LABELS = {
    "contract_company_name": "Компания контракта",
    "contract_company_slug": "Slug компании контракта",
    "contract_vessel_name": "Судно контракта",
    "contract_rank": "Должность в контракте",
    "home_airport": "Home airport",
    "departure_airport": "Departure airport",
    "departure_date": "Дата вылета",
}

SUBMISSION_LABELS = {
    "opening_vessel": "ПОДАЧА - opening m/v",
    "previous_vessel": "ПОДАЧА - previous m/v",
    "date_available_display": "ПОДАЧА - Date available display",
    "desirable_salary_display": "ПОДАЧА - Desirable salary display",
    "contract_duration_display": "ПОДАЧА - Contract duration display",
    "rank_since_sentence": "ПОДАЧА - фраза о стаже в должности",
    "coc_qr_paragraph": "ПОДАЧА - текст о QR codes на COC",
    "coc_gmdss_expiry_note": "ПОДАЧА - текст о сроке COC/GMDSS",
    "usa_visa_valid_paragraph": "ПОДАЧА - текст о сроке USA visa",
    "sb_expiry_paragraph": "ПОДАЧА - текст о сроке Seaman Book",
    "gent_name": "ПОДАЧА - имя кандидата для текста",
}


@dataclass(frozen=True)
class PlaceholderRow:
    section: str
    label: str
    placeholder: str


class Catalog:
    def __init__(self) -> None:
        self.rows: list[PlaceholderRow] = []
        self.seen: set[str] = set()

    def add(self, section: str, label: str, key_or_placeholder: str) -> None:
        placeholder = normalize_placeholder(key_or_placeholder)
        if placeholder in self.seen:
            return
        self.seen.add(placeholder)
        self.rows.append(PlaceholderRow(section, clean_label(label), placeholder))


def normalize_placeholder(value: str) -> str:
    text = str(value or "").strip()
    if text.startswith("{{") and text.endswith("}}"):
        inner = text[2:-2].strip()
        return "{{ " + inner + " }}"
    return "{{ " + text + " }}"


def clean_label(value: str) -> str:
    text = " ".join(str(value or "").replace("\n", " ").split())
    return text or "Поле"


def humanize_key(key: str) -> str:
    return key.replace("_", " ").strip()


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_dxa = int(sum(widths_inches) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = table._tbl.tblGrid
    if old_grid is not None:
        table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    table._tbl.insert(1, grid)

    for row in table.rows:
        for index, width in enumerate(widths_inches):
            cell = row.cells[index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 11.5, ACCENT_DARK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Manisov CRM - placeholders catalog")
    set_run_font(run, 8.5, MUTED)


def add_title(doc: Document, total: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manisov CRM")
    set_run_font(run, 24, ACCENT_DARK, True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Каталог плейсхолдеров Word-шаблонов")
    set_run_font(run2, 15, INK, True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(14)
    run3 = p3.add_run(f"Дата: {date.today().isoformat()} | Всего уникальных плейсхолдеров: {total}")
    set_run_font(run3, 10, MUTED)


def add_section_table(doc: Document, section: str, rows: list[PlaceholderRow]) -> None:
    doc.add_heading(section, level=1)
    table = doc.add_table(rows=1 + len(rows), cols=2)
    set_table_geometry(table, [2.45, 4.05])
    set_repeat_table_header(table.rows[0])
    headers = ["Название поля", "Плейсхолдер"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, 10, ACCENT_DARK, True)

    for row_index, row in enumerate(rows, start=1):
        name_cell = table.rows[row_index].cells[0]
        ph_cell = table.rows[row_index].cells[1]
        for cell, value, is_placeholder in (
            (name_cell, row.label, False),
            (ph_cell, row.placeholder, True),
        ):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            if is_placeholder:
                r.font.name = "Consolas"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                r.font.size = Pt(8.7)
                r.font.color.rgb = RGBColor.from_string("24292F")
            else:
                set_run_font(r, 9.2, INK)
    doc.add_paragraph()


def literal_string(node: ast.AST | None) -> str | None:
    return node.value if isinstance(node, ast.Constant) and isinstance(node.value, str) else None


def literal_string_tuple(node: ast.AST | None) -> tuple[str, ...]:
    if isinstance(node, (ast.Tuple, ast.List)):
        values = [literal_string(item) for item in node.elts]
        return tuple(item for item in values if item)
    value = literal_string(node)
    return (value,) if value else ()


def load_ast(path: Path) -> ast.Module:
    return ast.parse(path.read_text(encoding="utf-8"), filename=str(path))


def extract_specs_from_dict_constant(path: Path, names: set[str]) -> list[dict[str, object]]:
    tree = load_ast(path)
    specs: list[dict[str, object]] = []
    candidates: list[ast.AST] = []
    for node in ast.walk(tree):
        if isinstance(node, ast.Assign):
            target_names = {target.id for target in node.targets if isinstance(target, ast.Name)}
            if target_names.intersection(names):
                candidates.append(node.value)
        elif isinstance(node, ast.AnnAssign):
            target_name = node.target.id if isinstance(node.target, ast.Name) else None
            if target_name in names and node.value is not None:
                candidates.append(node.value)
    for value in candidates:
        values: Iterable[ast.AST]
        if isinstance(value, ast.Tuple):
            values = value.elts
        else:
            values = [value]
        for item in values:
            if not isinstance(item, ast.Dict):
                continue
            record: dict[str, object] = {}
            for key_node, value_node in zip(item.keys, item.values):
                key = literal_string(key_node)
                if not key:
                    continue
                if key == "legacy_prefixes":
                    record[key] = literal_string_tuple(value_node)
                else:
                    record[key] = literal_string(value_node)
            if record.get("placeholder_prefix"):
                specs.append(record)
    return specs


def extract_specs_from_spec_calls(path: Path) -> list[dict[str, object]]:
    tree = load_ast(path)
    specs: list[dict[str, object]] = []
    for node in ast.walk(tree):
        if not isinstance(node, ast.Call):
            continue
        if not isinstance(node.func, ast.Name) or node.func.id != "_spec":
            continue
        if len(node.args) < 2:
            continue
        code = literal_string(node.args[0])
        label = literal_string(node.args[1])
        if not code or not label:
            continue
        legacy: tuple[str, ...] = ()
        for kw in node.keywords:
            if kw.arg == "legacy_prefixes":
                legacy = literal_string_tuple(kw.value)
        prefix = re.sub(r"[^a-z0-9]+", "_", code.lower()).strip("_")
        specs.append(
            {
                "code": code,
                "certificate_type": label,
                "placeholder_prefix": prefix,
                "legacy_prefixes": legacy,
            }
        )
    return specs


def extract_schema_fields() -> dict[str, list[str]]:
    text = SCHEMA_PATH.read_text(encoding="utf-8")
    fields: dict[str, list[str]] = defaultdict(list)
    current: str | None = None
    for line in text.splitlines():
        class_match = re.match(r"class\s+(\w+)\(", line)
        if class_match:
            current = class_match.group(1)
            continue
        if current is None:
            continue
        field_match = re.match(r"\s{4}(\w+):\s*Mapped\[", line)
        if field_match and "relationship(" not in line:
            fields[current].append(field_match.group(1))
    return fields


def extract_tuple_pairs(path: Path, constant_name: str) -> list[tuple[str, str]]:
    tree = load_ast(path)
    pairs: list[tuple[str, str]] = []
    candidates: list[ast.AST] = []
    for node in ast.walk(tree):
        if isinstance(node, ast.Assign):
            if any(isinstance(target, ast.Name) and target.id == constant_name for target in node.targets):
                candidates.append(node.value)
        elif isinstance(node, ast.AnnAssign):
            if isinstance(node.target, ast.Name) and node.target.id == constant_name and node.value is not None:
                candidates.append(node.value)
    for value in candidates:
        if not isinstance(value, ast.Tuple):
            continue
        for item in value.elts:
            if isinstance(item, ast.Tuple) and len(item.elts) >= 2:
                key = literal_string(item.elts[0])
                label = literal_string(item.elts[1])
                if key and label:
                    pairs.append((key, label))
    return pairs


def extract_returned_dict_keys(path: Path, function_name: str, prefix: str) -> list[str]:
    tree = load_ast(path)
    keys: list[str] = []
    for node in ast.walk(tree):
        if not isinstance(node, ast.FunctionDef) or node.name != function_name:
            continue
        for child in ast.walk(node):
            if isinstance(child, ast.Dict):
                for key_node in child.keys:
                    key = literal_string(key_node)
                    if key and key.startswith(prefix):
                        keys.append(key)
    return sorted(set(keys))


def extract_js_field_defs(path: Path, array_name: str) -> list[tuple[str, str, str]]:
    text = path.read_text(encoding="utf-8")
    start = text.find(f"export const {array_name}")
    if start < 0:
        return []
    end = text.find("];", start)
    if end < 0:
        return []
    block = text[start:end]
    rows: list[tuple[str, str, str]] = []
    for match in re.finditer(
        r'key:\s*"([^"]+)".*?label:\s*"([^"]+)".*?placeholder:\s*"([^"]+)"',
        block,
        flags=re.S,
    ):
        rows.append((match.group(1), match.group(2), match.group(3)))
    return rows


def extract_context_keys_from_file(path: Path) -> set[str]:
    text = path.read_text(encoding="utf-8")
    keys = set(re.findall(r'context\["([A-Za-z_][A-Za-z0-9_]*)"\]', text))
    keys.update(re.findall(r"context\.setdefault\(\s*\"([A-Za-z_][A-Za-z0-9_]*)\"", text))
    return keys


def extract_assign_doc_prefixes(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8")
    return sorted(set(re.findall(r'_assign_doc_fields\(\s*context,\s*"([A-Za-z_][A-Za-z0-9_]*)"', text)))


def extract_docx_placeholders(path: Path) -> set[str]:
    placeholders: set[str] = set()
    try:
        with ZipFile(path) as zf:
            xml_names = [
                name
                for name in zf.namelist()
                if name.startswith("word/")
                and name.endswith(".xml")
                and (
                    name == "word/document.xml"
                    or name.startswith("word/header")
                    or name.startswith("word/footer")
                    or name.startswith("word/footnotes")
                    or name.startswith("word/endnotes")
                )
            ]
            chunks: list[str] = []
            for name in xml_names:
                try:
                    root = ET.fromstring(zf.read(name))
                except ET.ParseError:
                    continue
                for node in root.iter():
                    if node.tag.endswith("}t") and node.text:
                        chunks.append(node.text)
            text = "".join(chunks)
    except Exception:
        return placeholders
    for expr in re.findall(r"\{\{\s*(.*?)\s*\}\}", text, flags=re.S):
        cleaned = " ".join(expr.split())
        if cleaned:
            placeholders.add(cleaned)
    return placeholders


def add_slot_rows(
    catalog: Catalog,
    *,
    section: str,
    slot_label: str,
    prefix: str,
    suffixes: tuple[str, ...],
) -> None:
    for suffix in suffixes:
        label = f"{slot_label} / {SUFFIX_LABELS.get(suffix, humanize_key(suffix))}"
        catalog.add(section, label, f"{prefix}_{suffix}")


def add_legacy_rows(catalog: Catalog, section: str, slot_label: str, legacy_prefixes: Iterable[str], *, include_certificate_number: bool) -> None:
    for legacy_prefix in legacy_prefixes:
        if include_certificate_number:
            catalog.add(section, f"{slot_label} / legacy / Номер сертификата", f"{legacy_prefix}_certificate_number")
        for suffix in ("document_number", "issue_date", "expiry_date", "issuing_authority"):
            catalog.add(
                section,
                f"{slot_label} / legacy / {SUFFIX_LABELS.get(suffix, humanize_key(suffix))}",
                f"{legacy_prefix}_{suffix}",
            )


def build_catalog() -> Catalog:
    catalog = Catalog()

    # Ukrainian contract fields from frontend definitions.
    for _key, label, placeholder in extract_js_field_defs(
        PROJECT_ROOT / "app" / "frontend" / "src" / "ukrContractFields.js",
        "UKR_CONTRACT_FIELD_DEFS",
    ):
        catalog.add("Украинский контракт", label, placeholder)

    # Contract tab fields and personal fields.
    contract_path = PROJECT_ROOT / "app" / "contract_fields.py"
    for key, label in extract_tuple_pairs(contract_path, "CANDIDATE_PERSONAL_PLACEHOLDER_FIELDS"):
        catalog.add("Кандидат - основные поля", label, key)
    for key, label in extract_tuple_pairs(contract_path, "CONTRACT_EDITABLE_FIELDS"):
        catalog.add("Контракт", label, key)
    for key, label in extract_tuple_pairs(contract_path, "CONTRACT_DEPARTURE_FIELDS"):
        catalog.add("Контракт", label, key)
    for key, label in extract_tuple_pairs(contract_path, "CONTRACT_AIRPORT_FIELDS"):
        catalog.add("Контракт", label, key)
    for key, label in CONTRACT_BASE_LABELS.items():
        catalog.add("Контракт", label, key)

    # Salary placeholders.
    for key in extract_returned_dict_keys(
        PROJECT_ROOT / "app" / "salary_calculator.py",
        "salary_placeholders_from_saved",
        "salary_",
    ):
        catalog.add("Salary calculator", SALARY_LABELS.get(key, humanize_key(key)), key)

    # Contract vessel and dynamic company/vessel patterns.
    vessel_specs = extract_tuple_pairs(PROJECT_ROOT / "app" / "vessel_specs.py", "VESSEL_FIELD_SPECS")
    for field_key, label in vessel_specs:
        if field_key == "name":
            continue
        suffix = "type" if field_key == "vessel_type" else field_key
        catalog.add("Контракт - судно", f"Контрактное судно / {label}", f"contract_vessel_{suffix}")
        catalog.add("Компании и суда - динамический шаблон", f"Судно компании / {label}", f"company_<company_slug>_<vessel_slug>_{suffix}")

    # Candidate top-level fields and loop item fields from schema.
    schema_fields = extract_schema_fields()
    for field in schema_fields.get("Candidate", []):
        if field not in CANDIDATE_EXCLUDED_FIELDS:
            catalog.add("Кандидат - все поля карточки", f"Candidate / {humanize_key(field)}", field)
    for class_name, (section, alias) in MODEL_CLASSES.items():
        for field in schema_fields.get(class_name, []):
            catalog.add(section, f"{class_name} / {humanize_key(field)}", f"{alias}.{field}")

    # Legacy aliases from _augment_template_context and document-field assignment prefixes.
    main_path = PROJECT_ROOT / "app" / "main.py"
    for key in sorted(extract_context_keys_from_file(main_path)):
        if key.startswith("ukr_") or key.startswith("salary_") or key.startswith("contract_"):
            continue
        catalog.add("Alias и вычисляемые поля", f"Alias / {humanize_key(key)}", key)
    for prefix in extract_assign_doc_prefixes(main_path):
        add_slot_rows(
            catalog,
            section="Legacy document/certificate aliases",
            slot_label=f"Legacy {prefix}",
            prefix=prefix,
            suffixes=("document_number", "issue_date", "expiry_date", "issuing_authority"),
        )

    # Submission/PODACHA context.
    for key in sorted(extract_context_keys_from_file(PROJECT_ROOT / "app" / "submission_pack.py")):
        if key in SUBMISSION_LABELS:
            catalog.add("ПОДАЧА", SUBMISSION_LABELS[key], key)

    # Canonical documents.
    doc_specs = extract_specs_from_dict_constant(
        PROJECT_ROOT / "app" / "canonical_documents.py",
        {"CANONICAL_DOCUMENT_SPECS"},
    )
    for spec in doc_specs:
        slot_label = str(spec.get("document_type") or spec.get("code") or spec.get("placeholder_prefix"))
        add_slot_rows(
            catalog,
            section="Documents - канонические слоты",
            slot_label=slot_label,
            prefix=str(spec["placeholder_prefix"]),
            suffixes=("document_number", "issue_date", "expiry_date", "issuing_authority", "place_of_issue"),
        )

    # Canonical visas.
    visa_specs = extract_specs_from_dict_constant(
        PROJECT_ROOT / "app" / "canonical_visas.py",
        {"CANONICAL_VISA_SPECS"},
    )
    for spec in visa_specs:
        slot_label = str(spec.get("document_type") or spec.get("code") or spec.get("placeholder_prefix"))
        add_slot_rows(
            catalog,
            section="Visas - канонические слоты",
            slot_label=slot_label,
            prefix=str(spec["placeholder_prefix"]),
            suffixes=(
                "document_number",
                "issue_date",
                "expiry_date",
                "issuing_authority",
                "place_of_issue",
                "visa_code",
                "visa_name",
            ),
        )

    # Diplomas and tanker diplomas.
    diploma_specs = extract_specs_from_dict_constant(
        PROJECT_ROOT / "app" / "canonical_diplomas.py",
        {"CANONICAL_DIPLOMA_SPECS", "CANONICAL_TANKER_DIPLOMA_SPECS"},
    )
    for spec in diploma_specs:
        slot_label = str(spec.get("certificate_type") or spec.get("code") or spec.get("placeholder_prefix"))
        add_slot_rows(
            catalog,
            section="Diplomas - канонические слоты",
            slot_label=slot_label,
            prefix=str(spec["placeholder_prefix"]),
            suffixes=("certificate_number", "issue_date", "expiry_date", "issuing_authority", "country_of_issue"),
        )
        if spec.get("code") == "COC":
            catalog.add("Diplomas - канонические слоты", f"{slot_label} / {SUFFIX_LABELS['competency_rank']}", "coc_competency_rank")
        add_legacy_rows(
            catalog,
            "Diplomas - legacy",
            slot_label,
            spec.get("legacy_prefixes") or (),
            include_certificate_number=True,
        )

    # Medical.
    medical_specs = extract_specs_from_dict_constant(
        PROJECT_ROOT / "app" / "canonical_medical.py",
        {"CANONICAL_MEDICAL_SPECS"},
    )
    for spec in medical_specs:
        slot_label = str(spec.get("certificate_type") or spec.get("code") or spec.get("placeholder_prefix"))
        add_slot_rows(
            catalog,
            section="Medical - канонические слоты",
            slot_label=slot_label,
            prefix=str(spec["placeholder_prefix"]),
            suffixes=("certificate_number", "issue_date", "expiry_date", "issuing_authority", "country_of_issue"),
        )
        add_legacy_rows(
            catalog,
            "Medical - legacy",
            slot_label,
            spec.get("legacy_prefixes") or (),
            include_certificate_number=True,
        )

    # STCW, ECDIS, Company and BWTS certificates.
    cert_specs = extract_specs_from_spec_calls(PROJECT_ROOT / "app" / "canonical_certificates.py")
    for spec in cert_specs:
        slot_label = str(spec.get("certificate_type") or spec.get("code") or spec.get("placeholder_prefix"))
        add_slot_rows(
            catalog,
            section="Certificates - канонические слоты",
            slot_label=slot_label,
            prefix=str(spec["placeholder_prefix"]),
            suffixes=("certificate_number", "issue_date", "expiry_date", "issuing_authority", "country_of_issue"),
        )
        add_legacy_rows(
            catalog,
            "Certificates - legacy",
            slot_label,
            spec.get("legacy_prefixes") or (),
            include_certificate_number=False,
        )

    # Placeholders actually present in uploaded/built-in templates.
    for template_path in sorted((PROJECT_ROOT / "templates").rglob("*")):
        if template_path.suffix.lower() != ".docx":
            continue
        for expression in sorted(extract_docx_placeholders(template_path)):
            catalog.add(
                "Найдено в DOCX-шаблонах",
                f"Шаблон / {expression}",
                expression,
            )

    return catalog


def build_document(catalog: Catalog) -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc, len(catalog.rows))

    grouped: dict[str, list[PlaceholderRow]] = defaultdict(list)
    for row in catalog.rows:
        grouped[row.section].append(row)

    section_order = [
        "Кандидат - основные поля",
        "Кандидат - все поля карточки",
        "Alias и вычисляемые поля",
        "ПОДАЧА",
        "Контракт",
        "Контракт - судно",
        "Украинский контракт",
        "Salary calculator",
        "Компании и суда - динамический шаблон",
        "Documents - канонические слоты",
        "Visas - канонические слоты",
        "Diplomas - канонические слоты",
        "Diplomas - legacy",
        "Medical - канонические слоты",
        "Medical - legacy",
        "Certificates - канонические слоты",
        "Certificates - legacy",
        "Legacy document/certificate aliases",
        "Цикл applications",
        "Цикл documents",
        "Цикл certificates",
        "Цикл flag_documents",
        "Цикл sea_service",
        "Цикл family_contacts",
        "Цикл attachments",
        "Найдено в DOCX-шаблонах",
    ]

    for section in section_order:
        rows = grouped.get(section)
        if not rows:
            continue
        rows = sorted(rows, key=lambda item: (item.label.lower(), item.placeholder.lower()))
        add_section_table(doc, f"{section} ({len(rows)})", rows)

    remaining = sorted(set(grouped) - set(section_order))
    for section in remaining:
        rows = sorted(grouped[section], key=lambda item: (item.label.lower(), item.placeholder.lower()))
        add_section_table(doc, f"{section} ({len(rows)})", rows)

    return doc


def main() -> None:
    catalog = build_catalog()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document(catalog)
    try:
        doc.save(str(OUTPUT))
        print(f"Wrote {OUTPUT}")
        print(f"Unique placeholders: {len(catalog.rows)}")
    except PermissionError:
        alternate = OUTPUT.with_name(OUTPUT.stem + ".generated" + OUTPUT.suffix)
        doc.save(str(alternate))
        print(f"Wrote {alternate} because {OUTPUT.name} is locked")
        print(f"Unique placeholders: {len(catalog.rows)}")


if __name__ == "__main__":
    main()
