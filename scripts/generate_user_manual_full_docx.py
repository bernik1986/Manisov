"""Generate full end-user manual (Word). Run: python scripts/generate_user_manual_full_docx.py"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "docs" / "reports" / "USER_MANUAL_FULL_WORKFLOW_RU.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = val
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    title = doc.add_heading("CrewDeck CRM — полное руководство пользователя", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Версия документа: {date.today().isoformat()}\n")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub.add_run(
        "Система управления данными моряков: карточки кандидатов, документы, сертификаты, "
        "шаблоны Word, пакеты ПОДАЧА, компании и суда."
    ).font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, "Содержание (оглавление)", 1)
    toc = [
        "1. О системе и типовой рабочий день",
        "2. Вход, выход и навигация",
        "3. Роли и права доступа",
        "4. Dashboard (главное меню)",
        "5. Seamens Data — список кандидатов",
        "6. Карточка кандидата — обзор",
        "7. Профиль и заявка (recruitment)",
        "8. Documents, Diplomas, Медицина, Certificates",
        "9. Sea service, Flag documents, Family contacts",
        "10. Сканы: загрузка, имена файлов, сроки",
        "11. Сгенерировать документы (шаблоны DOCX)",
        "12. ПОДАЧА — ZIP-пакет для рассылки",
        "13. Украинский контракт",
        "14. Templates — менеджер шаблонов",
        "15. Company — компании и суда",
        "16. Notifications — уведомления",
        "17. Users и Logs (администратор)",
        "18. Типовые сценарии (workflow)",
        "19. Ошибки и поддержка",
    ]
    add_bullets(doc, toc)

    add_heading(doc, "1. О системе и типовой рабочий день", 1)
    add_para(
        doc,
        "CrewDeck CRM (в интерфейсе — Maritime Management) предназначена для рекрутингового офиса: "
        "вести базу моряков, хранить документы и сертификаты со сканами, готовить письма и пакеты "
        "для судовладельцев на основе шаблонов Word.",
    )
    add_heading(doc, "Типовой цикл работы рекрутера", 2)
    add_numbered(
        doc,
        [
            "Войти в систему → открыть Seamens Data.",
            "Загрузить анкету (DOC/DOCX/PDF/XLS) или создать пустую карточку.",
            "Проверить и дополнить данные в карточке; прикрепить сканы к Documents и Certificates.",
            "При необходимости сгенерировать отдельные DOCX или собрать ZIP через ПОДАЧА.",
            "Следить за Notifications (просрочки, отсутствие сканов).",
            "Администратор ведёт пользователей, шаблоны, компании и журнал действий.",
        ],
    )

    add_heading(doc, "2. Вход, выход и навигация", 1)
    add_heading(doc, "2.1. Вход", 2)
    add_numbered(
        doc,
        [
            "Откройте адрес CRM в браузере (у вашей организации свой URL; локально часто http://127.0.0.1:5173).",
            "Введите логин и пароль, нажмите «Войти».",
            "При успехе откроется Dashboard. При ошибке — проверьте данные или обратитесь к администратору.",
        ],
    )
    add_heading(doc, "2.2. Выход", 2)
    add_para(doc, "Внизу левого меню — кнопка «Выйти». Сеанс завершается; для продолжения работы нужен повторный вход.")
    add_heading(doc, "2.3. Боковое меню", 2)
    add_table(
        doc,
        ["Раздел", "Назначение"],
        [
            ["Dashboard", "Обзор, быстрые переходы"],
            ["Seamens Data", "Список кандидатов, загрузка анкет"],
            ["Templates", "Папки и файлы шаблонов DOCX/PDF"],
            ["Company", "Компании, суда, плейсхолдеры для шаблонов"],
            ["Notifications", "Просрочки и отсутствующие сканы"],
            ["Logs", "Журнал действий (только admin)"],
            ["Users", "Учётные записи (только admin)"],
        ],
    )
    add_para(doc, "На внутренних страницах доступна кнопка «В меню» для возврата на Dashboard.")

    add_heading(doc, "3. Роли и права доступа", 1)
    add_table(
        doc,
        ["Роль", "Просмотр", "Редактирование карточек", "ПОДАЧА / загрузка", "Админ-разделы"],
        [
            ["admin", "Да", "Да", "Да", "Users, Logs, удаление кандидата"],
            ["recruiter", "Да", "Да", "Да", "Нет"],
            ["viewer", "Да", "Нет", "Нет", "Нет"],
        ],
    )
    add_bullets(
        doc,
        [
            "Viewer не видит пункты Users и Logs; прямой переход на эти URL перенаправляет в меню.",
            "Удалить кандидата целиком может только admin.",
            "Редактирование документов, сертификатов, стажа, flag documents — admin и recruiter.",
        ],
    )

    add_heading(doc, "4. Dashboard (главное меню)", 1)
    add_para(
        doc,
        "После входа отображается Dashboard: краткая сводка (число кандидатов, превью уведомлений) "
        "и кнопки перехода в Seamens Data и Notifications.",
    )

    add_heading(doc, "5. Seamens Data — список кандидатов", 1)
    add_heading(doc, "5.1. Загрузка анкеты", 2)
    add_para(doc, "Форма «Анкета» принимает: DOC, DOCX, XLS, XLSX, PDF. Система разбирает файл и создаёт или обновляет карточку.")
    add_para(
        doc,
        "Если кандидат с такими же ключевыми данными уже есть, появится запрос: обновить существующую карточку данными из новой анкеты.",
    )
    add_heading(doc, "5.2. Загрузка CV", 2)
    add_para(doc, "Форма «CV» — только PDF. Извлекаются поля резюме в карточку.")
    add_heading(doc, "5.3. Новая пустая карточка", 2)
    add_para(doc, "Кнопка «Новая пустая карточка» создаёт запись без файла и сразу открывает профиль для ручного ввода.")
    add_heading(doc, "5.4. Поиск и фильтры", 2)
    add_bullets(
        doc,
        [
            "Строка поиска — по имени, фамилии и связанным полям.",
            "Должность — по полям Position / Rank Applied For в блоке заявки (не по морскому стажу).",
            "Флот / тип судна — по последнему контракту в Sea service (если пусто, в списке «-»).",
            "Фильтры сохраняются в адресной строке; при смене фильтра страница сбрасывается на первую.",
        ],
    )
    add_heading(doc, "5.5. Пагинация", 2)
    add_para(
        doc,
        "Внизу списка: первая / предыдущая / номера страниц / следующая / последняя. "
        "Подпись «Показано X–Y из Z» показывает текущий диапазон.",
    )
    add_heading(doc, "5.6. Открытие карточки", 2)
    add_para(doc, "Нажмите «Открыть карточку» в строке таблицы или перейдите по ссылке из уведомления.")

    add_heading(doc, "6. Карточка кандидата — обзор", 1)
    add_para(
        doc,
        "Карточка открывается по адресу /candidates/{id}. Вверху — панель действий; ниже — вкладки разделов. "
        "Одновременно открыт один раздел (кнопки-вкладки под панелью).",
    )
    add_heading(doc, "6.1. Панель действий", 2)
    add_table(
        doc,
        ["Кнопка", "Действие"],
        [
            ["ПОДАЧА", "Собрать ZIP: info-list DOCX + выбранные сканы"],
            ["Сгенерировать документы", "Скачать один или несколько DOCX по шаблонам"],
            ["Украинский контракт", "Ручное заполнение полей для укр. шаблона"],
            ["Удалить кандидата", "Только admin, с подтверждением"],
        ],
    )
    add_heading(doc, "6.2. Вкладки карточки", 2)
    add_bullets(
        doc,
        [
            "Группы профиля: основная карточка, персональные данные, семья, образование, профессиональные данные, медицинские summary, submission/info-list и др.",
            "Заявка / recruitment — данные заявки (должность, зарплата, proposed vessel и т.д.).",
            "Documents — паспорта, seaman book, визы и прочие документы (канонический список + свои типы).",
            "Diplomas — дипломы (COC, GMDSS, tanker и др.).",
            "Медицина — медицинские сертификаты (канонические слоты).",
            "Certificates — STCW и прочие сертификаты по категориям.",
            "Sea service — морской стаж (открывается в полноэкранном окне).",
            "Flag documents — документы по флагам государств.",
            "Family contacts — контакты родственников.",
        ],
    )

    add_heading(doc, "7. Профиль и заявка (recruitment)", 1)
    add_para(
        doc,
        "В каждой группе профиля после изменений нажмите «Сохранить» в этой группе. "
        "Поля submission/info-list (Home Airport, Desirable Salary, ECDIS, Leaving Reason и др.) "
        "используются в шаблонах info-list и ПОДАЧА.",
    )
    add_para(
        doc,
        "Вкладка «Заявка / recruitment»: должность, дата заявки, proposed vessel, зарплата и др. "
        "Для редактирования нужна роль admin или recruiter. Нажмите «Сохранить заявку» после правок.",
    )

    add_heading(doc, "8. Documents, Diplomas, Медицина, Certificates", 1)
    add_heading(doc, "8.1. Общие действия в таблицах", 2)
    add_bullets(
        doc,
        [
            "Добавить строку — форма внизу раздела (для канонических слотов часть полей уже задана).",
            "Редактировать / Сохранить / Отмена — в строке.",
            "Удалить запись — с подтверждением.",
            "Колонка «Номер» — номер документа/сертификата.",
            "Колонка «Код» — код слота (для имён сканов и плейсхолдеров).",
        ],
    )
    add_heading(doc, "8.2. Сроки действия (сертификаты и дипломы)", 2)
    add_bullets(
        doc,
        [
            "Режим «+5 лет» — вторая дата считается автоматически от даты выдачи или окончания.",
            "«Unlimited» — без даты окончания.",
            "«Другое» — обе даты вручную; кнопка «Применить +5 лет» пересчитывает недостающую дату.",
        ],
    )
    add_para(
        doc,
        "Строки с истекающим или просроченным сроком подсвечиваются. "
        "Дата окончания не может быть раньше даты выдачи — система покажет ошибку.",
    )

    add_heading(doc, "9. Sea service, Flag documents, Family contacts", 1)
    add_heading(doc, "9.1. Sea service", 2)
    add_para(
        doc,
        "Вкладка Sea service открывает полноэкранное окно: добавление, редактирование, удаление записей стажа "
        "(судно, тип, должность, даты, IMO и т.д.). Закрытие: кнопка «Закрыть», клик по фону или Escape.",
    )
    add_heading(doc, "9.2. Flag documents и Family contacts", 2)
    add_para(
        doc,
        "Таблицы с добавлением, редактированием и удалением (admin/recruiter). "
        "К каждой записи flag document можно прикрепить скан.",
    )

    add_heading(doc, "10. Сканы: загрузка, имена файлов, сроки", 1)
    add_heading(doc, "10.1. Загрузка скана", 2)
    add_numbered(
        doc,
        [
            "В строке Documents / Certificates / Diplomas / Медицина / Flag documents нажмите загрузку или перетащите PDF/изображение в зону строки.",
            "Изображения конвертируются в PDF.",
            "Скан можно скачать, заменить или удалить.",
        ],
    )
    add_heading(doc, "10.2. Имена файлов", 2)
    add_para(
        doc,
        "При сохранении и в ZIP ПОДАЧА имя формируется автоматически:",
    )
    add_para(doc, "{код должности} {Фамилия} {код слота}.pdf", bold=True)
    add_para(doc, "Пример: CO Ivanov AFF.pdf — должность CO, фамилия Ivanov, код сертификата AFF.")
    add_para(
        doc,
        "Код должности берётся из current rank или заявки. Код слота — из колонки «Код» в таблице. "
        "Заполните должность в карточке и код в строке до загрузки — иначе имя может быть неполным.",
    )

    add_heading(doc, "11. Сгенерировать документы (шаблоны DOCX)", 1)
    add_numbered(
        doc,
        [
            "На карточке нажмите «Сгенерировать документы».",
            "Слева выберите папку из Templates, справа отметьте один или несколько файлов .docx.",
            "Нажмите «Сгенерировать выбранные» — браузер скачает готовые файлы с подставленными данными кандидата.",
        ],
    )
    add_para(
        doc,
        "В шаблонах используются плейсхолдеры вида {{ surname }}, {{ company_slug_vessel_slug_imo }} и др. "
        "Список плейсхолдеров компаний/судов см. в разделе Company (копирование из карточки судна).",
    )

    add_heading(doc, "12. ПОДАЧА — ZIP-пакет для рассылки", 1)
    add_para(
        doc,
        "ПОДАЧА — основной способ собрать пакет для судовладельца: один ZIP с info-list и сканами. "
        "Отдельно от «Сгенерировать документы» (там только DOCX без архива и без сканов).",
    )
    add_heading(doc, "12.1. Пошагово", 2)
    add_numbered(
        doc,
        [
            "Откройте карточку кандидата → «ПОДАЧА».",
            "Укажите opening m/v (судно подачи) и при необходимости previous m/v (для ex-crew). Эти поля только для этого ZIP, в карточку не сохраняются.",
            "Отметьте шаблоны DOCX (по умолчанию часто выбраны «инфо лист…»).",
            "Отметьте сканы из Documents и Certificates (только строки, где скан уже загружен).",
            "Нажмите «Собрать ZIP» — скачается файл PODACHA_…zip.",
            "Отправьте архив клиенту почтой или мессенджером вручную (CRM письма не отправляет).",
        ],
    )
    add_heading(doc, "12.2. Что проверить перед ПОДАЧА", 2)
    add_bullets(
        doc,
        [
            "Фамилия, имя, должность; даты паспорта, COC, seaman book.",
            "Поля submission/info-list в профиле (аэропорт, зарплата, ECDIS, вакцинация, leaving reason).",
            "Сканы прикреплены к нужным строкам.",
            "Шаблоны info-list лежат в Templates (папка Podacha или аналог).",
        ],
    )
    add_table(
        doc,
        ["Сообщение", "Действие"],
        [
            ["Выберите хотя бы один шаблон или скан", "Поставьте галочки"],
            ["нет скана", "Загрузите скан в карточке, снова откройте ПОДАЧА"],
            ["API ПОДАЧА не найден (404)", "Сообщите IT — перезапуск backend"],
        ],
    )

    add_heading(doc, "13. Украинский контракт", 1)
    add_para(
        doc,
        "Кнопка «Украинский контракт» открывает форму полей (укр. ФИО, паспорт, адреса и т.д.). "
        "Данные сохраняются в карточке. В шаблонах Word используйте плейсхолдеры {{ ukr_surname }}, {{ ukr_full_name_ua }} и др. "
        "Доступно admin и recruiter.",
    )

    add_heading(doc, "14. Templates — менеджер шаблонов", 1)
    add_bullets(
        doc,
        [
            "Слева — дерево папок (корень Templates), справа — файлы выбранной папки.",
            "+ Root folder / + Subfolder — создать папки (admin/recruiter).",
            "Browse или Drag & Drop — загрузить .doc, .docx, .pdf.",
            "Переименование и удаление папок/файлов — с подтверждением.",
            "Скачать — загрузить файл на компьютер.",
            "Refresh — обновить список.",
            "Поиск по дереву — фильтр папок и файлов.",
        ],
    )
    add_para(
        doc,
        "Шаблоны из этого раздела доступны в «Сгенерировать документы» и в окне ПОДАЧА.",
    )

    add_heading(doc, "15. Company — компании и суда", 1)
    add_para(
        doc,
        "Справочник компаний и судов для плейсхолдеров в DOCX (например, {{ company_century_janina_imo }}).",
    )
    add_heading(doc, "15.1. Ручное ведение", 2)
    add_bullets(
        doc,
        [
            "+ Папка / + Подпапка — структура папок.",
            "+ Компания — компания в выбранной папке.",
            "Выберите компанию → справа список судов → + Судно (имя, IMO, флаг, тип).",
            "В карточке судна — копирование плейсхолдеров в буфер обмена.",
            "Refresh — обновить дерево.",
        ],
    )
    add_heading(doc, "15.2. Импорт из Excel", 2)
    add_para(
        doc,
        "Кнопка «Импорт Excel» (admin/recruiter): загрузите файл .xlsx/.xls с колонками:",
    )
    add_table(
        doc,
        ["Колонка", "Содержание"],
        [
            ["Company", "Название компании (при смене строки — новая компания)"],
            ["IMO", "Номер IMO судна"],
            ["Vessel name", "Название судна"],
        ],
    )
    add_para(
        doc,
        "Импорт идёт в выбранную папку (или в корень Companies). "
        "Существующие компании и суда не дублируются. После импорта появится сводка: создано / пропущено.",
    )

    add_heading(doc, "16. Notifications — уведомления", 1)
    add_para(
        doc,
        "Страница показывает непрочитанные уведомления по кандидатам: просроченные документы/сертификаты, "
        "скоро истекающие, отсутствие сканов.",
    )
    add_numbered(
        doc,
        [
            "Раскройте блок кандидата («Раскрыть (N)»).",
            "Нажмите текст уведомления — откроется карточка с подсветкой нужной строки документа или сертификата.",
        ],
    )

    add_heading(doc, "17. Users и Logs (администратор)", 1)
    add_heading(doc, "17.1. User Management", 2)
    add_bullets(
        doc,
        [
            "Создать пользователя: логин, пароль, полное имя, роль (viewer / recruiter / admin).",
            "Смена роли и пароля в строке пользователя.",
            "Active / Inactive — отключить доступ без удаления.",
            "Удаление — с подтверждением; свою учётку удалить нельзя.",
            "Нельзя снять роль admin с последнего активного администратора.",
        ],
    )
    add_heading(doc, "17.2. Audit Logs", 2)
    add_para(
        doc,
        "Журнал действий в системе: кто, когда, что изменил. Фильтры: пользователь, даты, кнопка «Применить». "
        "Быстрые периоды: 1 / 3 / 7 / 14 / 30 дней.",
    )

    add_heading(doc, "18. Типовые сценарии (workflow)", 1)
    add_heading(doc, "18.1. Новый кандидат с анкеты", 2)
    add_numbered(
        doc,
        [
            "Seamens Data → загрузить DOCX/PDF анкеты → открыть карточку.",
            "Проверить персональные данные и recruitment.",
            "Дозагрузить сканы в Documents и Certificates.",
            "ПОДАЧА → выбрать info-list и сканы → Собрать ZIP.",
        ],
    )
    add_heading(doc, "18.2. Обновление данных существующего кандидата", 2)
    add_numbered(
        doc,
        [
            "Загрузить анкету повторно → подтвердить обновление дубликата.",
            "Или вручную отредактировать поля и сохранить.",
            "Проверить Notifications на просрочки.",
        ],
    )
    add_heading(doc, "18.3. Подготовка шаблона для офиса", 2)
    add_numbered(
        doc,
        [
            "Templates → создать папку → загрузить DOCX с плейсхолдерами.",
            "На тестовой карточке — «Сгенерировать документы» для проверки.",
            "Для пакетов подачи — положить info-list в папку, используемую в ПОДАЧА.",
        ],
    )
    add_heading(doc, "18.4. Перенос справочника судов с Excel", 2)
    add_numbered(
        doc,
        [
            "Подготовить Excel (Company, IMO, Vessel name).",
            "Company → выбрать папку → Импорт Excel.",
            "Проверить дерево и плейсхолдеры на карточке судна.",
        ],
    )

    add_heading(doc, "19. Ошибки и поддержка", 1)
    add_numbered(
        doc,
        [
            "Проверьте вход в систему и роль (viewer не редактирует).",
            "Проверьте формат и размер файла (анкета до ~40 МБ, скан до ~15 МБ — лимиты могут отличаться на сервере).",
            "Обновите страницу (F5) и повторите действие.",
            "Скопируйте точный текст ошибки для IT или администратора.",
        ],
    )
    add_para(
        doc,
        "Сообщите поддержке: раздел, что нажали, текст ошибки (скриншот). "
        "Для ПОДАЧА и шаблонов укажите имя кандидата и имя файла шаблона.",
    )

    add_heading(doc, "Приложение: связанные документы", 1)
    add_bullets(
        doc,
        [
            "docs/PODACHA_INSTRUCTION_STAFF_RU.md — детальная инструкция по ПОДАЧА (есть Word-версия).",
            "docs/INTERFACE_GUIDE_USER_RU.md — краткая инструкция по интерфейсу.",
            "docs/SCAN_FILENAME_CONVENTION.md — коды должностей и слотов для имён сканов.",
            "docs/SEAMENS_DATA_FILTERS_RU.md — фильтры списка кандидатов.",
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"\n— Конец документа. CrewDeck CRM, {date.today().year} —")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    try:
        doc.save(str(OUTPUT))
        print(f"Wrote {OUTPUT}")
    except PermissionError:
        alt = OUTPUT.with_name(OUTPUT.stem + ".generated" + OUTPUT.suffix)
        doc.save(str(alt))
        print(f"Note: {OUTPUT} is locked; wrote {alt}")


if __name__ == "__main__":
    main()
