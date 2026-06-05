"""Generate stakeholder plan DOCX for Contract tab feature. Run: python scripts/generate_contract_tab_plan_docx.py"""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "docs" / "reports" / "CONTRACT_TAB_DEVELOPMENT_PLAN_RU.docx"
REPORTS_DIR = PROJECT_ROOT / "docs" / "reports"
ASSETS_CANDIDATES = [
    PROJECT_ROOT / "assets" / "contract_tab_mockup.png",
    PROJECT_ROOT / "assets" / "contract_generate_modal_mockup.png",
    Path(r"C:\Users\berni\.cursor\projects\c-Users-berni-Documents-Parcer\assets\contract_tab_mockup.png"),
    Path(r"C:\Users\berni\.cursor\projects\c-Users-berni-Documents-Parcer\assets\contract_generate_modal_mockup.png"),
]
IMG_TAB = REPORTS_DIR / "contract_tab_mockup.png"
IMG_MODAL = REPORTS_DIR / "contract_generate_modal_mockup.png"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = val
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, caption: str, width_in: float = 6.2) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()
    else:
        add_para(doc, f"[Макет: {caption} — файл не найден: {path}]", bold=True)


def resolve_mockup_images() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    pairs = [
        (ASSETS_CANDIDATES[0], ASSETS_CANDIDATES[2], IMG_TAB),
        (ASSETS_CANDIDATES[1], ASSETS_CANDIDATES[3], IMG_MODAL),
    ]
    for local, cursor_path, dest in pairs:
        for src in (local, cursor_path):
            if src.exists():
                shutil.copy2(src, dest)
                break


def build() -> Document:
    doc = Document()
    title = doc.add_heading("План разработки: вкладка «Контракт» в карточке кандидата", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        f"Документ для стейкхолдеров · CRM Parcer · {date.today().strftime('%d.%m.%Y')}\n"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub.add_run("Статус: планирование (код продукта не изменялся)").font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, "1. Цель и бизнес-задача", 1)
    add_para(
        doc,
        "Дать рекрутеру единое место в карточке кандидата для подготовки морского контракта: "
        "выбрать компанию-работодателя, судно и должность, увидеть автоматически подставленные "
        "реквизиты, связать расчёт зарплаты из калькулятора и сгенерировать DOCX только из "
        "заранее подготовленной папки шаблонов «Контракты».",
    )
    add_bullets(
        doc,
        [
            "Сократить время сборки контракта и снизить ошибки ручного копирования.",
            "Использовать уже заведённые в CRM компании, суда и матрицу зарплат.",
            "Обеспечить стабильные плейсхолдеры в Word-шаблонах для каждого поля вкладки.",
        ],
    )

    add_heading(doc, "2. Текущее состояние системы (аудит, без доработок)", 1)
    add_table(
        doc,
        ["Компонент", "Что уже есть", "Как используем в «Контракт»"],
        [
            [
                "Companies Manager",
                "Справочник компаний и судов (IMO, флаг, GRT, DWT, engine…), плейсхолдеры company_{slug}_{vessel}_*",
                "Дропдаун компаний и зависимый дропдаун судов",
            ],
            [
                "Калькулятор зарплаты",
                "Вкладка в карточке: company + rank → матрица Salary Scale; сохранение в salary_calculation_json; плейсхолдеры salary_*",
                "Подтягивать сохранённый расчёт; при смене company/rank — пересчёт или предупреждение",
            ],
            [
                "Должности",
                "CANONICAL_POSITION_OPTIONS (25 рангов) + матрица по компании из калькулятора",
                "Дропдаун должности (канон + ранги с матрицей для выбранной компании)",
            ],
            [
                "Генерация DOCX",
                "POST /candidates/{id}/generate/{template}; контекст из _serialize_candidate_context + плейсхолдеры",
                "Кнопка «Создать контракт» — только папка «Контракты»",
            ],
            [
                "Украинский трудовой договор",
                "Отдельный блок ukr_contract_json (~20 полей ukr_*)",
                "Не смешивать: «Контракт» — морской/COE; UA — отдельный сценарий при необходимости",
            ],
            [
                "Шаблоны",
                "Templates Manager — дерево папок и DOCX",
                "Создать папку «Контракты» (или «Contracts») и ограничить выбор ею",
            ],
        ],
    )

    add_heading(doc, "3. Предлагаемый UX (вкладка «Контракт»)", 1)
    add_para(doc, "Новая вкладка в навигации карточки кандидата (рядом с «Калькулятор зарплаты», «Визы»).", bold=True)

    add_heading(doc, "3.1. Блок выбора (каскадные дропдауны)", 2)
    add_table(
        doc,
        ["Поле", "Источник данных", "Поведение"],
        [
            ["Компания *", "GET /companies-manager → companies[]", "Сортировка по имени; обязательное"],
            [
                "Судно",
                "vessels[] где company_id = выбранная компания",
                "Disabled пока нет компании; опционально «— не выбрано —»",
            ],
            [
                "Должность *",
                "CANONICAL_POSITION_OPTIONS + ranks из матрицы зарплат компании",
                "При выборе — подгрузка salary template; синхрон с калькулятором",
            ],
        ],
    )

    add_heading(doc, "3.2. Автозаполнение под дропдаунами (только чтение + копирование)", 2)
    add_para(doc, "После выбора трёх параметров показываются карточки (не редактируются вручную — источник правды в справочниках):")
    add_bullets(
        doc,
        [
            "Компания: название, slug, при необходимости адрес/реквизиты (если добавим в Companies).",
            "Судно: название, IMO, флаг, Port of Registry, GRT, Deadweight, Year Built, Engine… (поля из Companies → Vessels).",
            "Должность: каноническое имя rank.",
            "Зарплата: блок из сохранённого salary_calculation_json, если company+rank совпадают; иначе подсказка «Сохраните расчёт в калькуляторе».",
            "Кандидат (кратко): ФИО, дата рождения, паспорт, SB — из карточки для превью контракта.",
        ],
    )

    add_heading(doc, "3.3. Дополнительные поля вкладки (редактируемые)", 2)
    add_para(
        doc,
        "Помимо автоподстановки — поля, специфичные для контракта (даты посадки, порт, срок контракта, "
        "номер контракта, валюта, примечания). Список согласуется со стейкхолдерами по реальным DOCX-шаблонам.",
    )
    add_table(
        doc,
        ["Поле (пример)", "Плейсхолдер", "Тип"],
        [
            ["Дата подписания контракта", "{{ contract_sign_date }}", "дата"],
            ["Срок контракта / Period of employment", "{{ contract_period }}", "текст"],
            ["Дата посадки (embarkation)", "{{ contract_embarkation_date }}", "дата"],
            ["Порт посадки", "{{ contract_embarkation_port }}", "текст"],
            ["Номер контракта", "{{ contract_number }}", "текст"],
            ["Примечания", "{{ contract_remarks }}", "текст"],
        ],
    )
    add_para(
        doc,
        "Полный реестр полей формируется после инвентаризации всех файлов в папке «Контракты» "
        "(сканирование {{ … }} в DOCX).",
    )

    add_heading(doc, "3.4. Кнопки действий", 2)
    add_bullets(
        doc,
        [
            "«Сохранить» — записать contract_json на кандидата (аналог ukr_contract_json / salary_calculation_json).",
            "«Создать контракт» — модальное окно выбора шаблонов только из папки «Контракты» → скачивание DOCX.",
            "Опционально: «Синхронизировать с калькулятором» — перенести company/rank/total wage в калькулятор.",
        ],
    )

    add_heading(doc, "4. Макеты интерфейса (визуализация)", 1)
    add_picture_if_exists(
        doc,
        IMG_TAB,
        "Рис. 1 — Вкладка «Контракт»: дропдауны, автозаполнение, кнопка создания",
    )
    add_picture_if_exists(
        doc,
        IMG_MODAL,
        "Рис. 2 — Диалог «Создать контракт»: только шаблоны из папки «Контракты»",
    )

    add_heading(doc, "5. Архитектура данных", 1)
    add_heading(doc, "5.1. Хранение", 2)
    add_para(
        doc,
        "Рекомендуется новое поле candidates.contract_json (TEXT, JSON) без отдельной таблицы на первом этапе:",
    )
    add_bullets(
        doc,
        [
            "contract_company_id, contract_company_name",
            "contract_vessel_id, contract_vessel_name",
            "contract_rank",
            "contract_sign_date, contract_period, contract_embarkation_date, contract_embarkation_port, contract_number, contract_remarks",
            "contract_updated_at, contract_updated_by",
        ],
    )

    add_heading(doc, "5.2. API (черновик)", 2)
    add_table(
        doc,
        ["Метод", "Назначение"],
        [
            ["GET /candidates/{id}/contract-context", "Компании, суда, ранги, snapshot судна/компании, salary snapshot"],
            ["PUT /candidates/{id}/contract", "Сохранение contract_json"],
            ["GET /templates-manager/contracts", "Список DOCX только из папки «Контракты» (по folder_id или имени)"],
            ["POST /candidates/{id}/generate-contract", "Генерация с merged context (кандидат + contract + salary + vessel placeholders)"],
        ],
    )

    add_heading(doc, "5.3. Контекст для docxtpl", 2)
    add_para(doc, "При генерации контракта в context добавляются (помимо существующих полей кандидата):")
    add_table(
        doc,
        ["Группа", "Примеры плейсхолдеров"],
        [
            ["Выбор вкладки", "contract_company_name, contract_vessel_name, contract_rank, contract_*"],
            ["Компания/судно (справочник)", "company_{slug}_{vessel}_name, _imo, _grt, … (уже реализованы)"],
            ["Зарплата", "salary_total_wage, salary_basic_monthly_wage, salary_owners_bonus, … (уже реализованы)"],
            ["Кандидат", "surname, first_name, passport_number, sea_service[0], … (уже в контексте)"],
        ],
    )
    add_para(
        doc,
        "Для каждого нового поля вкладки — обязательная строка в реестре плейсхолдеров (UI «скопировать» как в Companies/Visas).",
    )

    add_heading(doc, "6. Ограничение шаблонов папкой «Контракты»", 1)
    add_bullets(
        doc,
        [
            "В Templates Manager создаётся корневая или дочерняя папка с точным именем «Контракты» (согласовать регистр).",
            "Backend: resolve folder_id по имени; фильтр template_files по folder_id и подпапкам (опционально).",
            "Frontend: модал генерации не показывает дерево целиком — только эту папку.",
            "Валидация: при генерации контракта проверять, что template_file_id принадлежит папке «Контракты».",
        ],
    )

    add_heading(doc, "7. Связь с калькулятором зарплаты", 1)
    add_table(
        doc,
        ["Сценарий", "Поведение"],
        [
            [
                "Company + Rank совпадают с salary_calculation_json",
                "Показать все salary_* в превью; генерация подставляет актуальные суммы",
            ],
            [
                "Не совпадают",
                "Жёлтый баннер: «Сохраните расчёт в калькуляторе для этой компании и должности»",
            ],
            [
                "Пользователь меняет company на вкладке Контракт",
                "Опционально предложить обновить калькулятор одной кнопкой",
            ],
        ],
    )

    add_heading(doc, "8. Этапы разработки и оценка", 1)
    add_table(
        doc,
        ["Этап", "Содержание", "Оценка"],
        [
            ["0. Подготовка", "Инвентаризация DOCX в папке «Контракты», реестр плейсхолдеров", "1–2 дн"],
            ["1. Backend", "contract_json, миграция, contract-context API, merge в _serialize_candidate_context", "2–3 дн"],
            ["2. Frontend вкладка", "ContractSection, каскадные дропдауны, превью, сохранение", "3–4 дн"],
            ["3. Генерация", "Фильтр папки, модал, endpoint generate-contract, тесты", "2 дн"],
            ["4. QA / приёмка", "E2E, проверка 2–3 реальных шаблонов, документация для рекрутеров", "2 дн"],
            ["Итого", "", "~10–13 рабочих дней"],
        ],
    )

    add_heading(doc, "9. Риски и зависимости", 1)
    add_bullets(
        doc,
        [
            "Шаблоны контрактов различаются по компаниям — нужен единый словарь плейсхолдеров или несколько подпапок.",
            "Дублирование company/rank между калькулятором и контрактом — нужна явная синхронизация.",
            "Суда без заполненных полей — в контракте пустые плейсхолдеры; дисциплина ведения Companies.",
            "Старые шаблоны с другими именами переменных — миграция шаблонов или alias в backend.",
        ],
    )

    add_heading(doc, "10. Критерии приёмки (для стейкхолдеров)", 1)
    add_bullets(
        doc,
        [
            "В карточке есть вкладка «Контракт» с тремя дропдаунами.",
            "При выборе компании список судов только этой компании.",
            "Под дропдаунами отображаются данные компании, судна, должности и зарплаты.",
            "Сохранение выбора восстанавливается при повторном открытии кандидата.",
            "«Создать контракт» показывает только шаблоны из папки «Контракты».",
            "Сгенерированный DOCX содержит корректные ФИО, судно, зарплату без ручной правки в Word.",
            "Для каждого поля вкладки в UI доступен текст плейсхолдера для верстальщика шаблонов.",
        ],
    )

    add_heading(doc, "11. Рекомендации перед стартом разработки", 1)
    add_bullets(
        doc,
        [
            "Подготовить 2–3 эталонных DOCX в папке «Контракты» и список желаемых {{ переменных }}.",
            "Утвердить полный список редактируемых полей вкладки (помимо автоподстановки).",
            "Утвердить: должность только из канонического списка или свободный ввод.",
            "Провести демо текущих Companies + Salary Calculator для согласования ожиданий.",
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        f"Сформировано: scripts/generate_contract_tab_plan_docx.py · Parcer CRM · {date.today().isoformat()}"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def main() -> None:
    resolve_mockup_images()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    if IMG_TAB.exists():
        print(f"  Mockup 1: {IMG_TAB}")
    if IMG_MODAL.exists():
        print(f"  Mockup 2: {IMG_MODAL}")


if __name__ == "__main__":
    main()
