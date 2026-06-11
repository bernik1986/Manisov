from __future__ import annotations

import logging
import re
import subprocess
from datetime import date, datetime
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from docx import Document as DocxDocument
from docx.table import Table
from sqlalchemy import func
from sqlalchemy.orm import Session

from models.schema import (
    Application,
    Candidate,
    Certificate,
    Document,
    FamilyContact,
    FlagDocument,
    Notification,
    SeaService,
)
from parser.base import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parser for DOCX application forms with table-based extraction."""

    _SECTION_TABLE_HINTS: dict[str, set[str]] = {
        "personal_data": {
            "surname",
            "first name",
            "date of birth",
            "nationality",
            "email",
            "telephone no",
            "address",
            "marital status",
            "personal data",
            "applied for position of",
        },
        "documents": {
            "passport",
            "passport no",
            "document",
            "document number",
            "date of issue",
            "date of expiry",
            "issuing authority",
            "documents",
            "number",
        },
        "certificates": {
            "certificate",
            "certificate no",
            "certificate number",
            "stcw",
            "coc",
            "cop",
            "date issued",
            "expiry date",
            "part b qualifications",
            "yes",
            "no",
        },
        "sea_service": {
            "vessel name",
            "vessel type",
            "rank",
            "sign on",
            "sign off",
            "employer",
            "imo",
            "contract duration",
            "previous sea service",
            "ship s name",
            "date on",
            "date off",
        },
    }
    _SECTION_HEADER_OVERRIDES: dict[str, dict[str, str]] = {
        "documents": {
            "documents": "document_type",
            "document": "document_type",
            "number": "document_number",
            "document number": "document_number",
            "issuing authority": "issuing_authority",
            "date issued": "date_of_issue",
            "date of issue": "date_of_issue",
            "expiry date": "date_of_expiry",
            "date of expiry": "date_of_expiry",
        },
        "certificates": {
            "certificates": "certificate_type",
            "certificate": "certificate_type",
            "number": "certificate_number",
            "certificate number": "certificate_number",
            "issuing authority": "issuing_authority",
            "date issued": "date_issued",
            "date of issue": "date_issued",
            "expiry date": "expiry_date",
            "date of expiry": "expiry_date",
        },
        "sea_service": {
            "vessel": "vessel_name",
            "vessel name": "vessel_name",
            "ship s name": "vessel_name",
            "rank": "rank_on_vessel",
            "sign on": "sign_on_date",
            "sign off": "sign_off_date",
            "date on": "sign_on_date",
            "date off": "sign_off_date",
            "employer": "employer",
            "imo": "imo_number",
            "contract duration": "contract_duration",
            "vessel type": "vessel_type",
        },
    }

    def parse(self, file_path: str | Path) -> dict[str, Any]:
        try:
            path = Path(file_path)
            extension = self.detect_format(path)
            if extension not in {".docx", ".doc"}:
                raise ValueError("DocxParser supports only .docx/.doc files")

            if extension == ".doc":
                with TemporaryDirectory() as temp_dir:
                    converted = self._convert_doc_to_docx(path, Path(temp_dir))
                    doc = DocxDocument(str(converted))
                    return self.ensure_result_contract(self._parse_docx_document(doc))

            doc = DocxDocument(str(path))
            return self.ensure_result_contract(self._parse_docx_document(doc))
        except Exception as exc:
            logger.exception("DOCX parsing failed for '%s': %s", file_path, exc)
            raise

    def _parse_docx_document(self, doc: DocxDocument) -> dict[str, Any]:
        if self._is_century_layout(doc):
            return self._parse_century_document(doc)
        if self._is_overseas_layout(doc):
            return self._parse_overseas_document(doc)

        result = self.empty_result()

        result["personal_data"] = self._parse_personal_data_table(doc)
        result["documents"] = self._parse_documents_table(doc)
        result["certificates"] = self._parse_certificates_table(doc)
        result["sea_service"] = self._parse_sea_service_table(doc)
        result["applications"] = self._parse_applications_from_personal(doc)

        # CR-RT 05A fallback strategy
        if not result["personal_data"] or self._is_low_quality_personal_data(result["personal_data"]):
            result["personal_data"] = self._parse_personal_data_crrt(doc)
        if self._is_low_quality_sea_service(result["sea_service"]):
            crrt_sea_service = self._parse_sea_service_crrt(doc)
            if not self._is_low_quality_sea_service(crrt_sea_service):
                result["sea_service"] = crrt_sea_service

        return result

    def _is_overseas_layout(self, doc: DocxDocument) -> bool:
        if len(doc.tables) < 3:
            return False
        t0_text = " ".join(
            self.normalize_field_label(cell.text)
            for row in doc.tables[0].rows[:8]
            for cell in row.cells
            if cell.text and cell.text.strip()
        )
        t2_text = " ".join(
            self.normalize_field_label(cell.text)
            for row in doc.tables[2].rows[:6]
            for cell in row.cells
            if cell.text and cell.text.strip()
        )
        return "rank applied for" in t0_text and "seaman book" in t2_text and "passport" in t2_text

    def _parse_overseas_document(self, doc: DocxDocument) -> dict[str, Any]:
        result = self.empty_result()
        result["personal_data"] = self._parse_personal_data_overseas(doc)
        self._enrich_overseas_personal(result["personal_data"])
        result["family_contacts"] = self._derive_family_contacts_overseas(result["personal_data"])
        docs, certs = self._parse_documents_and_certificates_overseas(doc)
        result["documents"] = docs
        result["certificates"] = certs
        result["sea_service"] = self._parse_sea_service_overseas(doc)
        position = result["personal_data"].get("position_applied_for")
        rank = result["personal_data"].get("rank_applied_for") or position
        if position or rank:
            result["applications"] = [{"position_applied_for": position or rank, "rank_applied_for": rank or position}]
        return result

    def _parse_personal_data_overseas(self, doc: DocxDocument) -> dict[str, Any]:
        if not doc.tables:
            return {}
        table = doc.tables[0]
        personal: dict[str, Any] = {}
        for row in table.rows:
            tokens = self._dedupe_consecutive([cell.text.strip().replace("\n", " ") for cell in row.cells])
            if not tokens:
                continue
            normalized = [self.normalize_field_label(token) for token in tokens]
            if "sss no" in normalized:
                continue
            for idx, token in enumerate(tokens):
                norm = self.normalize_field_label(token.rstrip(":"))
                if norm == "date place of birth":
                    value = self._next_overseas_value(tokens, idx)
                    if value:
                        if "/" in value:
                            dob, pob = [p.strip() for p in value.split("/", 1)]
                            if dob:
                                personal["date_of_birth"] = dob
                            if pob:
                                personal["place_of_birth"] = pob
                        else:
                            personal["date_of_birth"] = value
                    continue
                if norm == "citizenship age":
                    value = self._next_overseas_value(tokens, idx)
                    if value:
                        if "/" in value:
                            nationality, age = [p.strip() for p in value.split("/", 1)]
                            if nationality:
                                personal["nationality"] = nationality
                                personal["citizenship"] = nationality
                            if age:
                                personal["age"] = age
                        else:
                            personal["nationality"] = value
                    continue

                mapped = self.map_to_canonical_field(norm)
                if mapped:
                    value = self._next_overseas_value(tokens, idx)
                    if value:
                        personal[mapped] = value
        if "position_applied_for" in personal and "rank_applied_for" not in personal:
            personal["rank_applied_for"] = personal["position_applied_for"]
        return personal

    def _next_overseas_value(self, tokens: list[str], label_idx: int) -> str | None:
        label = tokens[label_idx].strip()
        for candidate in tokens[label_idx + 1 :]:
            c = candidate.strip()
            if not c:
                continue
            norm = self.normalize_field_label(c.rstrip(":"))
            if c.lower() == label.lower():
                continue
            if self.map_to_canonical_field(norm):
                break
            if c.endswith(":"):
                continue
            return c
        return None

    @staticmethod
    def _has_cyrillic(text: str) -> bool:
        return bool(re.search(r"[\u0400-\u04FF]", text))

    @staticmethod
    def _build_full_name_from_parts(
        surname: str | None,
        first_name: str | None,
        middle_name: str | None,
    ) -> str | None:
        parts: list[str] = []
        for value in (surname, first_name, middle_name):
            cleaned = (value or "").strip()
            if not cleaned:
                continue
            if parts and cleaned.lower() == parts[-1].lower():
                continue
            parts.append(cleaned)
        if not parts:
            return None
        return " ".join(parts)

    @staticmethod
    def _parse_comma_address(address: str) -> dict[str, str | None]:
        """Split overseas permanent address into city / region / country when labels are absent."""
        parts = [segment.strip() for segment in address.split(",") if segment.strip()]
        if not parts:
            return {}
        country = parts[-1]
        city = None
        street_parts = parts[:-1]
        while street_parts and re.match(r"^(ap\.?|apt\.?|#)\s*\d", street_parts[-1], re.IGNORECASE):
            street_parts.pop()
        if street_parts and not re.match(r"^\d+$", street_parts[-1]):
            city = street_parts[-1]
        result: dict[str, str | None] = {"country": country[:100] if country else None}
        if city:
            result["city"] = city[:100]
            result["region"] = city[:100]
        return result

    def _enrich_overseas_personal(self, personal: dict[str, Any]) -> None:
        """Fill CRM fields missing from overseas table layout using parsed name/address/spouse data."""
        if not personal:
            return

        if not personal.get("primary_phone") and personal.get("telephone_no"):
            personal["primary_phone"] = personal["telephone_no"]

        latin = self._build_full_name_from_parts(
            personal.get("surname"),
            personal.get("first_name"),
            personal.get("middle_name"),
        )
        if latin and not personal.get("latin_full_name"):
            personal["latin_full_name"] = latin[:255]
        if latin and not personal.get("full_name"):
            personal["full_name"] = latin[:255]

        name_blob = " ".join(
            str(personal.get(key) or "")
            for key in ("surname", "first_name", "middle_name", "spouse_name", "father_name", "mother_name")
        )
        if self._has_cyrillic(name_blob):
            native = self._build_full_name_from_parts(
                personal.get("surname"),
                personal.get("first_name"),
                personal.get("middle_name"),
            )
            if native and not personal.get("native_full_name"):
                personal["native_full_name"] = native[:255]
        elif latin and not personal.get("native_full_name"):
            personal["native_full_name"] = latin[:255]

        spouse = (personal.get("spouse_name") or "").strip()
        if spouse and not personal.get("marital_status"):
            personal["marital_status"] = "Married"

        address = (personal.get("permanent_address") or "").strip()
        if address:
            parsed_addr = self._parse_comma_address(address)
            if parsed_addr.get("country") and not personal.get("country"):
                personal["country"] = parsed_addr["country"]
            if parsed_addr.get("city") and not personal.get("city"):
                personal["city"] = parsed_addr["city"]
            if parsed_addr.get("region") and not personal.get("region"):
                personal["region"] = parsed_addr["region"]

        pob = (personal.get("place_of_birth") or "").strip()
        if pob and not personal.get("country_of_birth"):
            personal["country_of_birth"] = pob[:100]
        if pob and not personal.get("country") and "," not in pob:
            personal["country"] = pob[:100]

    def _derive_family_contacts_overseas(self, personal: dict[str, Any]) -> list[dict[str, Any]]:
        """Map spouse (and beneficiary-style fields if present) into family_contacts rows."""
        rows: list[dict[str, Any]] = []
        if not personal:
            return rows

        spouse = (personal.get("spouse_name") or "").strip()
        if spouse:
            item: dict[str, Any] = {
                "full_name": spouse[:150],
                "relationship_to_candidate": "Spouse",
                "contact_type": "Spouse",
            }
            parts = spouse.split(None, 1)
            if len(parts) == 2:
                item["surname"] = parts[0][:100]
                item["first_name"] = parts[1][:100]
            elif len(parts) == 1:
                item["first_name"] = parts[0][:100]
            phone = (personal.get("mobile_phone") or personal.get("secondary_phone") or "").strip()
            if phone:
                item["phone"] = phone[:50]
            addr = (personal.get("permanent_address") or "").strip()
            if addr:
                item["address"] = addr[:255]
            rows.append({k: v for k, v in item.items() if v not in (None, "")})

        existing_names = {(r.get("full_name") or "").strip().lower() for r in rows}
        for row in self._derive_family_contacts_century(personal):
            fn = (row.get("full_name") or "").strip().lower()
            if not fn or fn in existing_names:
                continue
            rows.append(row)
            existing_names.add(fn)

        return rows

    def _parse_documents_and_certificates_overseas(self, doc: DocxDocument) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        if len(doc.tables) < 3:
            return [], []
        table = doc.tables[2]
        documents: list[dict[str, Any]] = []
        certificates: list[dict[str, Any]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(cells):
                continue
            title = cells[0].strip()
            norm_title = self.normalize_field_label(title)
            if not title:
                continue
            details = self._parse_overseas_doc_row_details(cells[1:])
            if "cert of competency" in norm_title:
                cert_item: dict[str, Any] = {"certificate_type": "Certificate of Competency"}
                if details.get("number"):
                    cert_item["certificate_number"] = details["number"]
                if details.get("issue_date"):
                    cert_item["date_issued"] = details["issue_date"]
                if details.get("expiry_date"):
                    cert_item["expiry_date"] = details["expiry_date"]
                certificates.append(cert_item)
                continue

            doc_item: dict[str, Any] = {"document_type": title}
            if details.get("number"):
                doc_item["document_number"] = details["number"]
            if details.get("issue_date"):
                doc_item["date_of_issue"] = details["issue_date"]
            if details.get("expiry_date"):
                doc_item["date_of_expiry"] = details["expiry_date"]
            if doc_item:
                documents.append(doc_item)
        return documents, certificates

    def _parse_overseas_doc_row_details(self, cells: list[str]) -> dict[str, str]:
        details: dict[str, str] = {}
        for raw in cells:
            text = raw.strip()
            if not text or ":" not in text:
                continue
            key, value = [part.strip() for part in text.split(":", 1)]
            norm_key = self.normalize_field_label(key)
            if norm_key in {"no", "number"}:
                details["number"] = value
            elif norm_key in {"issue date", "date of issue"}:
                details["issue_date"] = value
            elif norm_key in {"exp date", "expiry date", "date of expiry", "date of validity"}:
                details["expiry_date"] = value
        return details

    def _parse_sea_service_overseas(self, doc: DocxDocument) -> list[dict[str, Any]]:
        if len(doc.tables) < 2:
            return []
        table = doc.tables[1]
        if len(table.rows) < 2:
            return []
        headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
        normalized_headers = [self.normalize_field_label(h) for h in headers]

        idx_rank = self._first_matching_header_index(normalized_headers, {"rank"})
        idx_vessel = self._first_matching_header_index(normalized_headers, {"vessel"})
        idx_type = self._first_matching_header_index(normalized_headers, {"type of vessel", "type"})
        idx_flag = self._first_matching_header_index(normalized_headers, {"flag"})
        idx_year = self._first_matching_header_index(normalized_headers, {"year built", "built"})
        idx_dwt = self._first_matching_header_index(normalized_headers, {"dwt", "grt"})
        idx_engine = self._first_matching_header_index(
            normalized_headers, {"engine type model", "engine type", "engine"}
        )
        idx_power = self._first_matching_header_index(
            normalized_headers, {"b h p kw", "bhp kw", "kw", "bhp", "power"}
        )
        idx_principal = self._first_matching_header_index(normalized_headers, {"name of principal", "principal"})
        idx_manning = self._first_matching_header_index(normalized_headers, {"manning agent", "manning"})
        from_to_indexes = [i for i, h in enumerate(normalized_headers) if h.startswith("from") or "from to" in h]
        idx_sign_on = from_to_indexes[0] if len(from_to_indexes) >= 1 else None
        idx_sign_off = from_to_indexes[1] if len(from_to_indexes) >= 2 else None
        idx_remarks = self._first_matching_header_index(
            normalized_headers, {"reason of discharge", "cause of discharge", "cause of", "discharge"}
        )

        rows_data: list[dict[str, Any]] = []
        for row in table.rows[1:]:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(values):
                continue
            vessel_name = self._value_at_index(values, idx_vessel)
            if not vessel_name:
                continue
            item: dict[str, Any] = {
                "vessel_name": vessel_name,
                "vessel_type": self._value_at_index(values, idx_type),
                "flag": self._value_at_index(values, idx_flag),
                "year_built": self._value_at_index(values, idx_year),
                "dwt": self._value_at_index(values, idx_dwt),
                "main_engine": self._value_at_index(values, idx_engine),
                "engine_power": self._value_at_index(values, idx_power),
                "rank_on_vessel": self._value_at_index(values, idx_rank),
                "sign_on_date": self._value_at_index(values, idx_sign_on),
                "sign_off_date": self._value_at_index(values, idx_sign_off),
                "employer": self._value_at_index(values, idx_principal),
                "manning_agency": self._value_at_index(values, idx_manning),
                "remarks": self._value_at_index(values, idx_remarks),
            }
            rows_data.append({k: v for k, v in item.items() if v not in (None, "")})
        return rows_data

    def _parse_century_document(self, doc: DocxDocument) -> dict[str, Any]:
        result = self.empty_result()
        result["personal_data"] = self._parse_personal_data_century(doc)
        self._enrich_century_contacts_from_phones_row(self._find_century_main_table(doc), result["personal_data"])
        self._enrich_century_matrix_related_personal(doc, result["personal_data"])
        result["family_contacts"] = self._derive_family_contacts_century(result["personal_data"])
        result["documents"] = self._parse_documents_century(doc)
        result["certificates"] = self._parse_certificates_century(doc)
        result["sea_service"] = self._parse_sea_service_century(doc)
        result["flag_documents"] = self._parse_flag_documents_century(doc)

        position = result["personal_data"].get("position_applied_for")
        if position:
            result["applications"] = [{"position_applied_for": position, "rank_applied_for": position}]
        return result

    @staticmethod
    def _is_low_quality_personal_data(personal: dict[str, Any]) -> bool:
        if not personal:
            return True
        string_values = [str(v).strip() for v in personal.values() if isinstance(v, str)]
        if not string_values:
            return True
        label_like = sum(1 for value in string_values if value.endswith(":"))
        return (label_like / len(string_values)) >= 0.4

    @staticmethod
    def _is_low_quality_sea_service(items: list[dict[str, Any]]) -> bool:
        if not items:
            return True
        informative_rows = 0
        for item in items:
            if not isinstance(item, dict):
                continue
            has_vessel = bool(str(item.get("vessel_name") or "").strip())
            has_rank = bool(str(item.get("rank_on_vessel") or "").strip())
            has_dates = bool(str(item.get("sign_on_date") or "").strip()) or bool(str(item.get("sign_off_date") or "").strip())
            if has_vessel or (has_rank and has_dates):
                informative_rows += 1
        return informative_rows == 0

    def _is_century_layout(self, doc: DocxDocument) -> bool:
        has_personal = False
        has_documentation = False
        for table in doc.tables:
            text = " ".join(
                self.normalize_field_label(cell.text)
                for row in table.rows[:40]
                for cell in row.cells
                if cell.text and cell.text.strip()
            )
            if "seaman s personal details" in text:
                has_personal = True
            if "c documentation" in text or ("documents" in text and "date of expiry" in text):
                has_documentation = True
        return has_personal or has_documentation

    def _find_century_main_table(self, doc: DocxDocument) -> Table | None:
        for table in doc.tables:
            text = " ".join(
                self.normalize_field_label(cell.text)
                for row in table.rows[:4]
                for cell in row.cells
                if cell.text and cell.text.strip()
            )
            if "seaman s personal details" in text:
                return table
        return None

    @staticmethod
    def _century_row_is_sea_service_table_header(normalized_tokens: list[str]) -> bool:
        """True when row is the vessel/rank/sign-on column header row (end of personal block)."""
        joined = " ".join(t for t in normalized_tokens if t)
        return (
            "vessel" in joined
            and "type" in joined
            and "rank" in joined
            and ("sign on" in joined or "s on" in joined)
        )

    def _parse_personal_data_century(self, doc: DocxDocument) -> dict[str, Any]:
        table = self._find_century_main_table(doc)
        if not table:
            return {}

        personal: dict[str, Any] = {}
        for row in table.rows[1:]:
            tokens = self._dedupe_consecutive([cell.text.strip().replace("\n", " ") for cell in row.cells])
            if not tokens:
                continue
            normalized_tokens = [self.normalize_field_label(token) for token in tokens]
            if self._century_row_is_sea_service_table_header(normalized_tokens):
                break
            if any("previous sea service" in token for token in normalized_tokens):
                continue
            if normalized_tokens and normalized_tokens[0] in {"cert of competency", "ceritificate no"}:
                continue
            row_has_passport = any(token == "passport no" for token in normalized_tokens)

            i = 0
            while i < len(tokens):
                label = tokens[i].strip()
                norm_label = self.normalize_field_label(label.rstrip(":"))
                mapped = self.map_to_canonical_field(norm_label)
                if not mapped:
                    i += 1
                    continue
                if mapped in {"passport_issue_date", "passport_expiry_date", "passport_place_of_issue"} and not row_has_passport:
                    i += 1
                    continue

                j = i + 1
                while j < len(tokens):
                    candidate = tokens[j].strip()
                    norm_candidate = self.normalize_field_label(candidate.rstrip(":"))
                    if not candidate:
                        j += 1
                        continue
                    next_mapped = self.map_to_canonical_field(norm_candidate)
                    if next_mapped:
                        if next_mapped == mapped:
                            j += 1
                            continue
                        break
                    if candidate.lower() == label.lower():
                        j += 1
                        continue
                    if candidate.endswith(":"):
                        j += 1
                        continue
                    personal[mapped] = candidate
                    break
                i += 1

        return personal

    def _enrich_century_contacts_from_phones_row(self, table: Table | None, personal: dict[str, Any]) -> None:
        """Fill primary/secondary phone and email from the wide 'Phones / Mob / EMAIL' row (CR-RT 05A)."""
        if not table or not personal:
            return
        for row in table.rows[1:]:
            raw = [c.text.strip().replace("\n", " ") for c in row.cells]
            if not any(raw):
                continue
            lead = " ".join(self.normalize_field_label(x) for x in raw[:4])
            if "phone" not in lead and "mob" not in lead and "email" not in lead:
                continue
            blob = " ".join(raw)
            emails = []
            for m in re.finditer(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", blob):
                e = m.group(0)
                if e not in emails:
                    emails.append(e)
            phones = []
            for m in re.finditer(r"\+?\d[\d\-\s()]{6,}\d", blob):
                p = re.sub(r"\s+", " ", m.group(0).strip())
                if p not in phones:
                    phones.append(p)
            if len(phones) >= 1:
                personal["primary_phone"] = phones[0]
            if len(phones) >= 2:
                personal["secondary_phone"] = phones[1]
            if emails:
                personal["email"] = emails[0]
            if len(emails) >= 2:
                personal["secondary_email"] = emails[1]
            return

    def _enrich_century_matrix_related_personal(self, doc: DocxDocument, personal: dict[str, Any]) -> None:
        """B(ii) Matrix block (CR-RT 05A) carries rank / vessel-type / watch / total sea lines outside the main personal table."""
        if personal is None:
            return
        for table in doc.tables:
            head_blob = " ".join(
                self.normalize_field_label(c.text)
                for r in table.rows[:4]
                for c in r.cells
                if c.text and c.text.strip()
            )
            if "matrix related information" not in head_blob:
                continue
            for row in table.rows[1:]:
                raw = [c.text.strip().replace("\n", " ") for c in row.cells]
                tokens = self._dedupe_consecutive(raw)
                if not tokens:
                    continue
                row_joined = " ".join(self.normalize_field_label(t) for t in tokens)
                if "c documentation" in row_joined:
                    break
                label = tokens[0].strip()
                norm_label = self.normalize_field_label(label.rstrip(":"))
                mapped = self.map_to_canonical_field(norm_label)
                if not mapped:
                    continue
                value: str | None = None
                j = 1
                while j < len(tokens):
                    candidate = tokens[j].strip()
                    if not candidate:
                        j += 1
                        continue
                    norm_candidate = self.normalize_field_label(candidate.rstrip(":"))
                    next_mapped = self.map_to_canonical_field(norm_candidate)
                    if next_mapped:
                        if next_mapped == mapped:
                            j += 1
                            continue
                        break
                    if candidate.lower() == label.lower():
                        j += 1
                        continue
                    if candidate.endswith(":") and len(candidate.strip()) <= 2:
                        j += 1
                        continue
                    value = candidate
                    break
                if value is None and len(tokens) >= 2:
                    tail = tokens[-1].strip()
                    if tail:
                        tail_norm = self.normalize_field_label(tail.rstrip(":"))
                        if not self.map_to_canonical_field(tail_norm):
                            value = tail
                if not value:
                    continue
                if mapped == "total_years_of_sea_service":
                    personal["total_sea_service"] = value.strip()
                personal[mapped] = value.strip()
            break

    def _derive_family_contacts_century(self, personal: dict[str, Any]) -> list[dict[str, Any]]:
        """Map beneficiary / next of kin from personal_data into family_contacts rows (Century forms)."""
        rows: list[dict[str, Any]] = []
        if not personal:
            return rows

        ben = (personal.get("beneficiary_full_name") or "").strip()
        if ben:
            addr = (personal.get("beneficiary_address") or personal.get("permanent_address") or "").strip()
            rel = (personal.get("beneficiary_relationship") or "").strip()
            item: dict[str, Any] = {
                "full_name": ben[:150],
                "relationship_to_candidate": (rel[:100] if rel else None),
                "address": (addr[:255] if addr else None),
            }
            bp = (personal.get("beneficiary_phone") or "").strip()
            if bp:
                item["phone"] = bp[:50]
            be = (personal.get("beneficiary_email") or "").strip()
            if be and "@" in be:
                item["email"] = be[:255]
            rows.append({k: v for k, v in item.items() if v not in (None, "")})

        nok = (personal.get("next_of_kin_full_name") or "").strip()
        if nok and nok.lower() != ben.lower():
            item2: dict[str, Any] = {"full_name": nok[:150]}
            r2 = (personal.get("next_of_kin_relationship") or "").strip()
            if r2:
                item2["relationship_to_candidate"] = r2[:100]
            a2 = (personal.get("next_of_kin_address") or "").strip()
            if a2:
                item2["address"] = a2[:255]
            p2 = (personal.get("next_of_kin_phone") or "").strip()
            if p2:
                item2["phone"] = p2[:50]
            rows.append({k: v for k, v in item2.items() if v not in (None, "")})

        return rows

    def _parse_documents_century(self, doc: DocxDocument) -> list[dict[str, Any]]:
        table = self._find_century_main_table(doc)
        if not table:
            return []

        start_labels = {
            "seaman s book no",
            "passport no",
            "usa visa no",
            "yellow fever vaccination",
        }
        rows_data: list[dict[str, Any]] = []
        for row in table.rows:
            tokens = self._dedupe_consecutive([cell.text.strip().replace("\n", " ") for cell in row.cells])
            if not tokens:
                continue
            norm_tokens = [self.normalize_field_label(token) for token in tokens]
            first_label = next((token for token in norm_tokens if token), "")
            if first_label not in start_labels:
                continue

            raw = " | ".join(tokens)
            item: dict[str, Any] = {}
            if "seaman s book" in first_label:
                item["document_type"] = "Seaman's Book"
            elif "passport" in first_label:
                item["document_type"] = "Passport"
            elif "usa visa" in first_label:
                item["document_type"] = "USA Visa"
            elif "yellow fever" in first_label:
                item["document_type"] = "Yellow Fever Vaccination"

            if len(tokens) > 1:
                for token in tokens[1:]:
                    norm = self.normalize_field_label(token)
                    if norm in {"date of issue", "date of validity", "place of issue"}:
                        continue
                    if self.map_to_canonical_field(norm):
                        continue
                    item["document_number"] = token
                    break

            date_issue = self._extract_century_value_after_label(tokens, "date of issue")
            if date_issue:
                item["date_of_issue"] = date_issue
            date_validity = self._extract_century_value_after_label(tokens, "date of validity")
            if date_validity:
                item["date_of_expiry"] = date_validity
            place_issue = self._extract_century_value_after_label(tokens, "place of issue")
            if place_issue:
                item["issuing_authority"] = place_issue

            if item:
                rows_data.append(item)

        return rows_data

    def _parse_certificates_century(self, doc: DocxDocument) -> list[dict[str, Any]]:
        table: Table | None = None
        for t in doc.tables:
            text = " ".join(
                self.normalize_field_label(cell.text)
                for row in t.rows[:20]
                for cell in row.cells
                if cell.text and cell.text.strip()
            )
            if "documentation" in text and "documents" in text:
                table = t
                break
        if not table:
            return []

        rows_data: list[dict[str, Any]] = []
        in_docs_section = False
        idx_cert: int | None = None
        idx_authority: int | None = None
        idx_number: int | None = None
        idx_issued: int | None = None
        idx_expiry: int | None = None
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(values):
                continue
            first = values[0].strip()
            norm_first = self.normalize_field_label(first)
            if not first:
                continue
            row_text = " ".join(self.normalize_field_label(v) for v in values if v)
            if "documents" in row_text and "date of expiry" in row_text:
                in_docs_section = True
                headers = [self.normalize_field_label(v) for v in values]
                idx_cert = self._first_matching_header_index(headers, {"documents", "document"})
                idx_authority = self._first_matching_header_index(headers, {"issuing authority"})
                idx_number = self._first_matching_header_index(headers, {"doc no", "document number", "number"})
                idx_issued = self._first_matching_header_index(headers, {"date of issuance", "date issued", "date of issue"})
                idx_expiry = self._first_matching_header_index(headers, {"date of expiry", "expiry date", "date of validity"})
                continue
            if "security checks" in row_text:
                break
            if not in_docs_section:
                continue
            if norm_first in {
                "fire fighting",
                "sso",
                "ssa",
                "safety officer",
            }:
                continue
            if norm_first in {"documents", "issuing authority"}:
                continue

            cert_name = self._value_at_index(values, idx_cert) or first
            item: dict[str, Any] = {"certificate_type": cert_name}
            authority = self._value_at_index(values, idx_authority) or ""
            if authority:
                item["issuing_authority"] = authority
            number = self._value_at_index(values, idx_number) or ""
            if number:
                item["certificate_number"] = number
            issued = self._value_at_index(values, idx_issued) or ""
            if issued:
                item["date_issued"] = issued
            expiry = self._value_at_index(values, idx_expiry) or ""
            if expiry:
                item["expiry_date"] = expiry

            rows_data.append(item)
        return rows_data

    def _extract_century_value_after_label(self, tokens: list[str], label: str) -> str | None:
        normalized = [self.normalize_field_label(token) for token in tokens]
        for idx, token in enumerate(normalized):
            if token != label:
                continue
            for candidate in tokens[idx + 1 :]:
                c = candidate.strip()
                if not c:
                    continue
                norm_c = self.normalize_field_label(c)
                if norm_c == label:
                    continue
                if self.map_to_canonical_field(norm_c):
                    continue
                return c
        return None

    def _parse_sea_service_century(self, doc: DocxDocument) -> list[dict[str, Any]]:
        ordered: list[Table] = []
        main = self._find_century_main_table(doc)
        if main is not None:
            ordered.append(main)
        for t in doc.tables:
            if t is not main:
                ordered.append(t)
        for table in ordered:
            rows = self._parse_sea_service_for_century_table(table)
            if rows:
                return rows
        return []

    def _parse_sea_service_for_century_table(self, table: Table) -> list[dict[str, Any]]:
        start_idx: int | None = None
        for idx, row in enumerate(table.rows):
            text = " ".join(self.normalize_field_label(cell.text) for cell in row.cells if cell.text and cell.text.strip())
            if "previous sea service" in text or "b i previous sea service" in text:
                start_idx = idx
                break
        if start_idx is None:
            return []

        header_values: list[str] = []
        normalized_headers: list[str] = []
        header_idx = start_idx + 1
        for delta in range(5):
            hi = start_idx + 1 + delta
            if hi >= len(table.rows):
                break
            hv = [cell.text.strip().replace("\n", " ") for cell in table.rows[hi].cells]
            nh = [self.normalize_field_label(v) if v.strip() else "" for v in hv]
            if self._first_matching_header_index(nh, {"vessel"}) is not None:
                header_values = hv
                normalized_headers = nh
                header_idx = hi
                break
        if not header_values:
            return []

        idx_vessel = self._first_matching_header_index(normalized_headers, {"vessel"})
        idx_type = self._first_matching_header_index(normalized_headers, {"type", "type of vessel"})
        idx_flag = self._first_matching_header_index(normalized_headers, {"flag"})
        idx_year = self._first_matching_header_index(normalized_headers, {"year built", "built"})
        idx_dwt = self._first_matching_header_index(normalized_headers, {"dwt", "grt"})
        idx_engine = self._first_matching_header_index(
            normalized_headers, {"engine type model", "engine type", "engine"}
        )
        idx_power = self._first_matching_header_index(
            normalized_headers, {"b h p kw", "bhp kw", "kw", "bhp", "power"}
        )
        idx_rank = self._first_matching_header_index(normalized_headers, {"rank"})
        idx_sign_on = self._first_matching_header_index(normalized_headers, {"sign on", "s on", "from"})
        idx_sign_off = self._first_matching_header_index(normalized_headers, {"sign off", "s off", "to"})
        idx_owner = self._first_matching_header_index(normalized_headers, {"owner", "principal"})
        idx_remarks = self._first_matching_header_index(
            normalized_headers, {"cause of discharge", "reason of discharge", "cause of", "discharge"}
        )

        rows_data: list[dict[str, Any]] = []
        for row in table.rows[header_idx + 1 :]:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if not any(values):
                continue
            if any("(dd/mm/yy)" in value.lower() for value in values if value):
                continue

            vessel_name = self._value_at_index(values, idx_vessel)
            norm_vessel = self.normalize_field_label(vessel_name or "")
            if not vessel_name:
                continue
            if "shipyard experience" in norm_vessel:
                break
            if norm_vessel in {"vessel", "type"}:
                continue

            item: dict[str, Any] = {
                "vessel_name": vessel_name,
                "vessel_type": self._value_at_index(values, idx_type),
                "flag": self._value_at_index(values, idx_flag),
                "year_built": self._value_at_index(values, idx_year),
                "main_engine": self._value_at_index(values, idx_engine),
                "engine_power": self._value_at_index(values, idx_power),
                "dwt": self._value_at_index(values, idx_dwt),
                "rank_on_vessel": self._value_at_index(values, idx_rank),
                "sign_on_date": self._value_at_index(values, idx_sign_on),
                "sign_off_date": self._value_at_index(values, idx_sign_off),
                "employer": self._value_at_index(values, idx_owner),
                "remarks": self._value_at_index(values, idx_remarks),
            }
            rows_data.append({k: v for k, v in item.items() if v not in (None, "")})
        return rows_data

    def _parse_flag_documents_century(self, doc: DocxDocument) -> list[dict[str, Any]]:
        """Extract D. FLAG DOCUMENTS table (CR-RT 05A / Century Bulker layout)."""
        rows_data: list[dict[str, Any]] = []
        for t_idx, table in enumerate(doc.tables):
            header_row_idx: int | None = None
            idx_flag = idx_rank = idx_doc = idx_issuance = idx_expiry = None
            for r_idx, row in enumerate(table.rows):
                vals = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if not any(vals):
                    continue
                norms = [self.normalize_field_label(v) for v in vals]
                i_flag = self._first_matching_header_index(norms, {"flag"})
                i_doc = self._first_matching_header_index(
                    norms, {"doc no", "doc number", "doc no ref", "document no", "document number"}
                )
                i_iss = self._first_matching_header_index(
                    norms,
                    {"date of issuance", "date of issue", "issue date", "issuance date"},
                )
                i_exp = self._first_matching_header_index(
                    norms,
                    {"date of expiry", "expiry date", "exp date"},
                )
                if i_flag is None or i_doc is None or i_iss is None or i_exp is None:
                    continue

                idx_flag, idx_rank, idx_doc = i_flag, self._first_matching_header_index(norms, {"rank"}), i_doc
                idx_issuance, idx_expiry = i_iss, i_exp
                header_row_idx = r_idx

                break

            if header_row_idx is None or idx_flag is None or idx_doc is None or idx_issuance is None or idx_expiry is None:
                continue

            for data_row in table.rows[header_row_idx + 1 :]:
                dvals = [cell.text.strip().replace("\n", " ") for cell in data_row.cells]
                if not any(v.strip() for v in dvals):
                    continue

                gate = " ".join(self.normalize_field_label(v) for v in dvals if v.strip())
                if "e security checks" in gate:
                    break
                first_nonempty = next((v.strip() for v in dvals if v.strip()), "")
                nf_first = self.normalize_field_label(first_nonempty)
                if (nf_first.startswith("e ") and "security" in gate) or (
                    nf_first.startswith("f ") and "questions" in gate and "flag documents" not in gate
                ):
                    break

                flag_val = self._value_at_index(dvals, idx_flag)
                if not flag_val:
                    continue
                if self.normalize_field_label(flag_val) == "flag":
                    continue

                rank_val = self._value_at_index(dvals, idx_rank)
                doc_no = self._value_at_index(dvals, idx_doc)
                issue_d = self._value_at_index(dvals, idx_issuance)
                exp_d = self._value_at_index(dvals, idx_expiry)

                item: dict[str, Any] = {
                    "flag_country": flag_val.strip()[:100],
                    "doc_number": doc_no.strip()[:100] if doc_no else None,
                    "date_of_issuance": issue_d.strip() if issue_d else None,
                    "date_of_expiry": exp_d.strip() if exp_d else None,
                }
                if rank_val:
                    item["rank"] = rank_val.strip()[:100]

                rows_data.append({k: v for k, v in item.items() if v not in (None, "")})

        return rows_data

    @staticmethod
    def _first_matching_header_index(headers: list[str], candidates: set[str]) -> int | None:
        """Match header by exact label or substring (longest candidate first)."""
        ordered = sorted(candidates, key=len, reverse=True)
        for cand in ordered:
            for idx, header in enumerate(headers):
                if not header:
                    continue
                if header == cand or cand in header or header in cand:
                    return idx
        return None

    @staticmethod
    def _value_at_index(values: list[str], idx: int | None) -> str | None:
        if idx is None:
            return None
        if idx >= len(values):
            return None
        value = values[idx].strip()
        return value or None

    @staticmethod
    def _dedupe_consecutive(values: list[str]) -> list[str]:
        out: list[str] = []
        prev: str | None = None
        for value in values:
            v = value.strip()
            if not v:
                continue
            if v == prev:
                continue
            out.append(v)
            prev = v
        return out

    @staticmethod
    def _is_label(text: str) -> bool:
        return text.endswith(":")

    def _parse_personal_data_crrt(self, doc: DocxDocument) -> dict[str, Any]:
        for table in doc.tables:
            if not table.rows:
                continue
            first_rows = " ".join(self.normalize_field_label(c.text) for r in table.rows[:2] for c in r.cells)
            if "personal data" not in first_rows:
                continue

            personal: dict[str, Any] = {}
            next_of_kin_mode = False
            for row in table.rows[2:]:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                vals = self._dedupe_consecutive(cells)
                if not vals:
                    continue

                if any("next of kin" in self.normalize_field_label(v) for v in vals):
                    next_of_kin_mode = True
                if any("native language" in self.normalize_field_label(v) for v in vals):
                    next_of_kin_mode = False

                i = 0
                while i < len(vals) - 1:
                    label = vals[i]
                    value = vals[i + 1]
                    if not self._is_label(label):
                        i += 1
                        continue
                    norm_label = self.normalize_field_label(label.rstrip(":"))
                    mapped = self.map_to_canonical_field(norm_label)

                    # Disambiguate "Surname/Name" inside Next of kin block.
                    if next_of_kin_mode and norm_label == "surname":
                        mapped = "next_of_kin_surname"
                    elif next_of_kin_mode and norm_label in {"name", "forename", "first name", "given name"}:
                        mapped = "next_of_kin_first_name"
                    elif next_of_kin_mode and norm_label in {"tel", "telephone number", "mobile"}:
                        mapped = "next_of_kin_phone"
                    elif norm_label == "name":
                        mapped = "first_name" if "first_name" not in personal else None
                    elif norm_label == "mobile":
                        mapped = "mobile_phone"
                    elif norm_label in {"tel", "telephone number"}:
                        mapped = "primary_phone"

                    if mapped and value and not self._is_label(value):
                        personal[mapped] = value
                    i += 1

            if personal:
                return personal
        return {}

    def _parse_applications_from_personal(self, doc: DocxDocument) -> list[dict[str, Any]]:
        for table in doc.tables:
            if not table.rows:
                continue
            first_rows = " ".join(self.normalize_field_label(c.text) for r in table.rows[:3] for c in r.cells)
            if "applied for position of" not in first_rows:
                continue
            for row in table.rows[2:6]:
                vals = self._dedupe_consecutive([cell.text.strip().replace("\n", " ") for cell in row.cells])
                for idx, token in enumerate(vals):
                    if "applied for position of" in self.normalize_field_label(token) and idx + 1 < len(vals):
                        value = vals[idx + 1]
                        if value:
                            return [{"position_applied_for": value, "rank_applied_for": value}]
        return []

    def _parse_sea_service_crrt(self, doc: DocxDocument) -> list[dict[str, Any]]:
        for table in doc.tables:
            if len(table.rows) < 4:
                continue
            top = " ".join(self.normalize_field_label(c.text) for c in table.rows[0].cells)
            if "previous sea service" not in top:
                continue

            header_vals = [cell.text.strip().replace("\n", " ") for cell in table.rows[2].cells]
            header_norm = [self.normalize_field_label(v) if v.strip() else "" for v in header_vals]
            idx_vessel = self._first_matching_header_index(header_norm, {"vessel", "ship"})
            idx_type = self._first_matching_header_index(header_norm, {"type", "vessel type"})
            idx_engine = self._first_matching_header_index(header_norm, {"engine"})
            idx_dwt = self._first_matching_header_index(header_norm, {"dwt", "grt"})
            idx_employer = self._first_matching_header_index(header_norm, {"owner", "employer", "principal"})
            idx_rank = self._first_matching_header_index(header_norm, {"rank"})
            idx_sign_on = self._first_matching_header_index(header_norm, {"sign on", "s on", "from"})
            idx_sign_off = self._first_matching_header_index(header_norm, {"sign off", "s off", "to"})
            idx_remarks = self._first_matching_header_index(
                header_norm, {"cause of discharge", "reason of discharge", "cause of", "discharge"}
            )
            idx_flag = self._first_matching_header_index(header_norm, {"flag"})
            idx_power = self._first_matching_header_index(header_norm, {"b h p kw", "bhp kw", "kw", "bhp"})

            items: list[dict[str, Any]] = []
            for row in table.rows[3:]:
                vals = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if len(vals) < 8 or not any(vals):
                    continue
                if "ship" in self.normalize_field_label(vals[0]) and "type" in self.normalize_field_label(vals[1]):
                    continue
                item = {
                    "vessel_name": self._value_at_index(vals, idx_vessel) or (vals[0] if idx_vessel is None else None),
                    "vessel_type": self._value_at_index(vals, idx_type) or (vals[1] if idx_type is None else None),
                    "main_engine": self._value_at_index(vals, idx_engine) or (vals[2] if idx_engine is None else None),
                    "dwt": self._value_at_index(vals, idx_dwt) or (vals[3] if idx_dwt is None else None),
                    "employer": self._value_at_index(vals, idx_employer) or (vals[4] if idx_employer is None else None),
                    "rank_on_vessel": self._value_at_index(vals, idx_rank) or (vals[5] if idx_rank is None else None),
                    "sign_on_date": self._value_at_index(vals, idx_sign_on) or (vals[6] if idx_sign_on is None else None),
                    "sign_off_date": self._value_at_index(vals, idx_sign_off) or (vals[7] if idx_sign_off is None else None),
                    "remarks": self._value_at_index(vals, idx_remarks) or (vals[8] if len(vals) > 8 and idx_remarks is None else None),
                    "flag": self._value_at_index(vals, idx_flag),
                    "engine_power": self._value_at_index(vals, idx_power),
                }
                vessel_name = (item.get("vessel_name") or "").strip()
                vessel_name_norm = self.normalize_field_label(vessel_name)
                if "shipyard experience" in vessel_name_norm or "sts experience" in vessel_name_norm:
                    break
                if vessel_name:
                    items.append(item)
            if items:
                return items
        return []

    def _convert_doc_to_docx(self, source_path: Path, output_dir: Path) -> Path:
        command = [
            "soffice",
            "--headless",
            "--convert-to",
            "docx",
            str(source_path),
            "--outdir",
            str(output_dir),
        ]
        try:
            subprocess.run(command, check=True, capture_output=True, text=True)
        except FileNotFoundError as exc:
            raise RuntimeError("soffice is not installed or not available in PATH") from exc
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or "").strip()
            raise RuntimeError(f"Failed to convert .doc to .docx: {stderr or exc}") from exc

        converted = output_dir / f"{source_path.stem}.docx"
        if not converted.exists():
            raise RuntimeError("DOC conversion finished, but .docx output was not found")
        return converted

    def _parse_personal_data_table(self, doc: DocxDocument) -> dict[str, Any]:
        table = self._find_table_by_section(doc.tables, "personal_data")
        if not table:
            return {}

        raw_data: dict[str, Any] = {}
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) < 2:
                continue
            key, value = cells[0], cells[1]
            raw_data[key] = value
        return self.map_fields(raw_data)

    def _parse_documents_table(self, doc: DocxDocument) -> list[dict[str, Any]]:
        table = self._find_table_by_section(doc.tables, "documents")
        if not table:
            return []
        return self._parse_grid_table(table, section="documents")

    def _parse_certificates_table(self, doc: DocxDocument) -> list[dict[str, Any]]:
        table = self._find_table_by_section(doc.tables, "certificates")
        if not table:
            return []
        return self._parse_grid_table(table, section="certificates")

    def _parse_sea_service_table(self, doc: DocxDocument) -> list[dict[str, Any]]:
        table = self._find_table_by_section(doc.tables, "sea_service")
        if not table:
            return []
        return self._parse_grid_table(table)

    def _find_table_by_section(self, tables: list[Table], section: str) -> Table | None:
        hints = self._SECTION_TABLE_HINTS.get(section, set())
        best_table: Table | None = None
        best_score = 0

        for table in tables:
            headers = self._extract_table_headers(table)
            if not headers:
                continue
            score = sum(1 for header in headers if header in hints)
            if score > best_score:
                best_score = score
                best_table = table

        # Require at least one section-specific header match.
        return best_table if best_score > 0 else None

    def _extract_table_headers(self, table: Table) -> set[str]:
        if not table.rows:
            return set()
        headers: set[str] = set()
        for row in table.rows[:3]:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    headers.add(self.normalize_field_label(cell.text))
        return headers

    def _parse_grid_table(self, table: Table, section: str | None = None) -> list[dict[str, Any]]:
        if not table.rows:
            return []

        section_overrides = self._SECTION_HEADER_OVERRIDES.get(section or "", {})
        header_row_idx = self._pick_header_row_index(table, section_overrides)
        headers = [cell.text.strip() for cell in table.rows[header_row_idx].cells]
        if not any(headers):
            return []

        canonical_headers = self._map_headers(headers, section_overrides)
        if not any(canonical_headers):
            return []
        rows_data: list[dict[str, Any]] = []

        for row in table.rows[header_row_idx + 1 :]:
            values = [cell.text.strip() for cell in row.cells]
            if not any(values):
                continue
            item: dict[str, Any] = {}
            for idx, value in enumerate(values):
                if not value:
                    continue
                canonical = canonical_headers[idx] if idx < len(canonical_headers) else None
                if canonical:
                    item[canonical] = value
            if item:
                rows_data.append(item)
        return rows_data

    def _map_headers(self, headers: list[str], section_overrides: dict[str, str]) -> list[str | None]:
        mapped: list[str | None] = []
        for header in headers:
            if not header:
                mapped.append(None)
                continue
            normalized = self.normalize_field_label(header)
            if normalized in section_overrides:
                mapped.append(section_overrides[normalized])
            else:
                mapped.append(self.map_to_canonical_field(header))
        return mapped

    def _pick_header_row_index(self, table: Table, section_overrides: dict[str, str]) -> int:
        best_idx = 0
        best_score = -1
        for idx, row in enumerate(table.rows[:3]):
            headers = [cell.text.strip() for cell in row.cells]
            mapped = self._map_headers(headers, section_overrides)
            score = sum(1 for item in mapped if item)
            if score > best_score:
                best_score = score
                best_idx = idx
        return best_idx

    @staticmethod
    def _to_date(value: Any) -> date | None:
        if value is None or value == "":
            return None
        if isinstance(value, date):
            return value
        if isinstance(value, str):
            raw = value.strip()
            if not raw:
                return None
            if raw.lower() in {"-", "--", "n/a", "na", "none", "unlimited", "no expiry", "without expiry"}:
                return None

            cleaned = re.sub(r"/\.", "/", raw)
            cleaned = re.sub(r"\s+", " ", cleaned).strip()
            for fmt in (
                "%Y-%m-%d",
                "%d.%m.%Y",
                "%d/%m/%Y",
                "%d-%m-%Y",
                "%d/%b/%Y",
                "%d/%B/%Y",
                "%d-%b-%Y",
                "%d-%B-%Y",
            ):
                try:
                    return datetime.strptime(cleaned, fmt).date()
                except ValueError:
                    continue
            return None
        return None

    @staticmethod
    def _to_float(value: Any) -> float | None:
        if value is None:
            return None
        if isinstance(value, (int, float)):
            return float(value)
        if isinstance(value, str):
            raw = value.strip()
            if not raw or raw.lower() in {"-", "--", "n/a", "na", "none"}:
                return None
            # Common merged-cell patterns in forms: "5137/115050", "75 000", etc.
            normalized = raw.replace(",", ".")
            match = re.search(r"-?\d+(?:\.\d+)?", normalized)
            if not match:
                return None
            try:
                return float(match.group(0))
            except ValueError:
                return None
        return None

    @staticmethod
    def _to_int(value: Any) -> int | None:
        parsed = DocxParser._to_float(value)
        if parsed is None:
            return None
        try:
            return int(parsed)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _to_bool(value: Any) -> bool | None:
        if value in (None, ""):
            return None
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return bool(value)
        if isinstance(value, str):
            normalized = value.strip().lower()
            if normalized in {"true", "1", "yes", "y"}:
                return True
            if normalized in {"false", "0", "no", "n"}:
                return False
        return None

    def _map_and_save_to_db(self, data: dict[str, Any], session: Session) -> Candidate:
        """
        Upsert candidate and replace related rows with fresh parsed data.

        Candidate match priority:
        1) candidate_id
        2) surname + date_of_birth
        """
        try:
            personal = data.get("personal_data", {}) or {}
            documents = data.get("documents", []) or []
            certificates = data.get("certificates", []) or []
            sea_service = data.get("sea_service", []) or []
            applications = data.get("applications", []) or []
            flag_documents = data.get("flag_documents", []) or []
            family_contacts = data.get("family_contacts", []) or []

            candidate: Candidate | None = None
            candidate_id = personal.get("candidate_id")
            if candidate_id not in (None, ""):
                try:
                    candidate = session.get(Candidate, int(candidate_id))
                except (TypeError, ValueError):
                    candidate = None

            if candidate is None:
                surname = personal.get("surname")
                dob = self._to_date(personal.get("date_of_birth"))
                if surname and dob:
                    # In real DB there may be duplicates for same surname+dob.
                    # Reuse the latest record instead of failing import.
                    candidate = (
                        session.query(Candidate)
                        .filter(
                            func.lower(Candidate.surname) == str(surname).strip().lower(),
                            Candidate.date_of_birth == dob,
                        )
                        .order_by(Candidate.candidate_id.desc())
                        .first()
                    )

            if candidate is None:
                candidate = Candidate()
                session.add(candidate)

            # Update candidate fields from parsed personal data.
            candidate_columns = {column.name for column in Candidate.__table__.columns}
            candidate_column_types = {
                column.name: column.type.__class__.__name__ for column in Candidate.__table__.columns
            }
            for key, value in personal.items():
                if key not in candidate_columns or key == "candidate_id":
                    continue
                column_type = candidate_column_types.get(key, "")
                if column_type == "Date":
                    coerced_value = self._to_date(value)
                elif column_type == "Integer":
                    coerced_value = self._to_int(value)
                elif column_type == "Float":
                    coerced_value = self._to_float(value)
                elif column_type == "Boolean":
                    coerced_value = self._to_bool(value)
                else:
                    coerced_value = value
                setattr(candidate, key, coerced_value)

            session.flush()

            # Replace related entities to remove stale records.
            existing_doc_ids = [
                row.document_id
                for row in session.query(Document.document_id)
                .filter(Document.candidate_id == candidate.candidate_id)
                .all()
            ]
            if existing_doc_ids:
                session.query(Notification).filter(Notification.document_id.in_(existing_doc_ids)).delete(
                    synchronize_session=False
                )
            existing_cert_ids = [
                row.certificate_id
                for row in session.query(Certificate.certificate_id)
                .filter(Certificate.candidate_id == candidate.candidate_id)
                .all()
            ]
            if existing_cert_ids:
                session.query(Notification).filter(Notification.certificate_id.in_(existing_cert_ids)).delete(
                    synchronize_session=False
                )
            session.query(Document).filter(Document.candidate_id == candidate.candidate_id).delete(synchronize_session=False)
            session.query(Certificate).filter(Certificate.candidate_id == candidate.candidate_id).delete(synchronize_session=False)
            session.query(SeaService).filter(SeaService.candidate_id == candidate.candidate_id).delete(synchronize_session=False)
            session.query(Application).filter(Application.candidate_id == candidate.candidate_id).delete(synchronize_session=False)
            session.query(FlagDocument).filter(FlagDocument.candidate_id == candidate.candidate_id).delete(synchronize_session=False)
            session.query(FamilyContact).filter(FamilyContact.candidate_id == candidate.candidate_id).delete(synchronize_session=False)

            doc_columns = {column.name for column in Document.__table__.columns}
            cert_columns = {column.name for column in Certificate.__table__.columns}
            service_columns = {column.name for column in SeaService.__table__.columns}
            app_columns = {column.name for column in Application.__table__.columns}
            flag_columns = {column.name for column in FlagDocument.__table__.columns}
            fam_columns = {column.name for column in FamilyContact.__table__.columns}

            for raw in documents:
                payload = {k: v for k, v in raw.items() if k in doc_columns}
                payload["candidate_id"] = candidate.candidate_id
                if "document_type" not in payload:
                    payload["document_type"] = raw.get("document_type") or "Unknown document"
                for date_key in ("date_of_issue", "date_of_expiry"):
                    if date_key in payload:
                        payload[date_key] = self._to_date(payload[date_key])
                session.add(Document(**payload))

            for raw in certificates:
                payload = {k: v for k, v in raw.items() if k in cert_columns}
                payload["candidate_id"] = candidate.candidate_id
                if "certificate_type" not in payload:
                    payload["certificate_type"] = raw.get("certificate_type") or "Unknown certificate"
                for date_key in ("date_issued", "expiry_date"):
                    if date_key in payload:
                        payload[date_key] = self._to_date(payload[date_key])
                session.add(Certificate(**payload))

            for raw in sea_service:
                payload = {k: v for k, v in raw.items() if k in service_columns}
                payload["candidate_id"] = candidate.candidate_id
                for date_key in ("sign_on_date", "sign_off_date"):
                    if date_key in payload:
                        payload[date_key] = self._to_date(payload[date_key])
                if "year_built" in payload:
                    payload["year_built"] = self._to_int(payload["year_built"])
                for float_key in ("dwt", "grt"):
                    if float_key in payload:
                        payload[float_key] = self._to_float(payload[float_key])
                session.add(SeaService(**payload))

            for raw in applications:
                payload = {k: v for k, v in raw.items() if k in app_columns}
                payload["candidate_id"] = candidate.candidate_id
                if not payload.get("position_applied_for") and not payload.get("rank_applied_for"):
                    continue
                for date_key in ("date_applied", "date_available"):
                    if date_key in payload:
                        payload[date_key] = self._to_date(payload[date_key])
                session.add(Application(**payload))

            for raw in flag_documents:
                payload = {k: v for k, v in raw.items() if k in flag_columns}
                payload["candidate_id"] = candidate.candidate_id
                if not payload.get("flag_country"):
                    continue
                for date_key in ("date_of_issuance", "date_of_expiry"):
                    if date_key in payload:
                        payload[date_key] = self._to_date(payload[date_key])
                session.add(FlagDocument(**payload))

            for raw in family_contacts:
                payload = {k: v for k, v in raw.items() if k in fam_columns}
                payload["candidate_id"] = candidate.candidate_id
                if not payload.get("full_name"):
                    full_name = raw.get("full_name") or " ".join(
                        part for part in (raw.get("first_name"), raw.get("surname")) if part
                    )
                    payload["full_name"] = full_name or "Unknown contact"
                session.add(FamilyContact(**payload))

            session.commit()
            session.refresh(candidate)
            return candidate
        except Exception as exc:
            session.rollback()
            logger.exception("Saving parsed DOCX payload failed: %s", exc)
            raise
