from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, event
from sqlalchemy.engine import Engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.main import app, get_db_session, pwd_context
from app.canonical_documents import CANONICAL_DOCUMENT_SPECS
from models.db import Base
from models.schema import Candidate, Document, Notification, Role, User
from tests.canonical_test_helpers import count_certificate_rows, find_certificate, find_document


@event.listens_for(Engine, "connect")
def _set_sqlite_pragma(dbapi_connection, _connection_record) -> None:
    try:
        cursor = dbapi_connection.cursor()
        cursor.execute("PRAGMA foreign_keys=ON")
        cursor.close()
    except Exception:
        pass


def _make_sample_docx(path: Path) -> Path:
    doc = DocxDocument()

    personal = doc.add_table(rows=4, cols=2)
    personal.cell(0, 0).text = "Surname"
    personal.cell(0, 1).text = "Petrov"
    personal.cell(1, 0).text = "First Name"
    personal.cell(1, 1).text = "Ivan"
    personal.cell(2, 0).text = "Date of Birth"
    personal.cell(2, 1).text = "1992-02-10"
    personal.cell(3, 0).text = "EMAIL"
    personal.cell(3, 1).text = "ivan.petrov@example.com"

    documents = doc.add_table(rows=2, cols=4)
    documents.cell(0, 0).text = "Document Type"
    documents.cell(0, 1).text = "Document Number"
    documents.cell(0, 2).text = "Date of Issue"
    documents.cell(0, 3).text = "Date of Expiry"
    documents.cell(1, 0).text = "Passport"
    documents.cell(1, 1).text = "AB123456"
    documents.cell(1, 2).text = "2020-01-01"
    documents.cell(1, 3).text = "2030-01-01"

    certificates = doc.add_table(rows=2, cols=3)
    certificates.cell(0, 0).text = "Certificate Type"
    certificates.cell(0, 1).text = "Certificate Number"
    certificates.cell(0, 2).text = "Expiry Date"
    certificates.cell(1, 0).text = "COC"
    certificates.cell(1, 1).text = "COC-7788"
    certificates.cell(1, 2).text = "2029-12-31"

    sea_service = doc.add_table(rows=2, cols=3)
    sea_service.cell(0, 0).text = "Vessel name"
    sea_service.cell(0, 1).text = "Rank"
    sea_service.cell(0, 2).text = "Sign on"
    sea_service.cell(1, 0).text = "SEA STAR"
    sea_service.cell(1, 1).text = "2/O"
    sea_service.cell(1, 2).text = "2023-06-01"

    doc.save(path)
    return path


def test_upload_then_get_candidate_in_memory_db(tmp_path: Path) -> None:
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    Base.metadata.create_all(bind=engine)

    def override_get_db_session():
        session = TestingSessionLocal()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides[get_db_session] = override_get_db_session

    try:
        with TestingSessionLocal() as session:
            admin_role = Role(name="admin", description="Full access")
            session.add(admin_role)
            session.flush()
            session.add(
                User(
                    username="admin",
                    password_hash=pwd_context.hash("admin123"),
                    full_name="Test Admin",
                    role_id=admin_role.role_id,
                    is_active=True,
                )
            )
            session.commit()

        client = TestClient(app)
        login_response = client.post(
            "/auth/login",
            json={"username": "admin", "password": "admin123"},
        )
        assert login_response.status_code == 200
        token = login_response.json()["access_token"]
        auth_headers = {"Authorization": f"Bearer {token}"}

        sample_file = _make_sample_docx(tmp_path / "sample_upload.docx")

        with sample_file.open("rb") as fh:
            upload_response = client.post(
                "/upload",
                files={
                    "file": (
                        "sample_upload.docx",
                        fh,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=auth_headers,
            )
        assert upload_response.status_code == 200
        upload_payload = upload_response.json()
        candidate_id = upload_payload["candidate_id"]
        assert upload_payload["duplicate"] is False

        candidate_response = client.get(f"/candidates/{candidate_id}", headers=auth_headers)
        assert candidate_response.status_code == 200
        payload = candidate_response.json()

        candidate = payload["candidate"]
        assert candidate["surname"] == "Petrov"
        assert candidate["first_name"] == "Ivan"
        assert candidate["email"] == "ivan.petrov@example.com"

        assert len(payload["documents"]) >= len(CANONICAL_DOCUMENT_SPECS)
        passport = find_document(
            payload["documents"],
            lambda row: (row.get("document_number") or "") == "AB123456",
        )
        assert passport is not None
        assert passport.get("document_number") == "AB123456"

        assert count_certificate_rows(payload) >= 1
        _section, coc = find_certificate(
            payload,
            lambda row: (row.get("certificate_number") or "") == "COC-7788"
            or (row.get("certificate_type") or "").upper() == "COC",
        )
        assert coc is not None

        with sample_file.open("rb") as fh:
            duplicate_response = client.post(
                "/upload",
                files={
                    "file": (
                        "sample_upload.docx",
                        fh,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=auth_headers,
            )
        assert duplicate_response.status_code == 200
        duplicate_payload = duplicate_response.json()
        assert duplicate_payload["duplicate"] is True
        assert duplicate_payload["candidate_id"] == candidate_id
        assert duplicate_payload["requires_confirmation"] is True
        assert duplicate_payload["updated"] is False
        assert "хотите обновить" in duplicate_payload["message"].lower()

        with sample_file.open("rb") as fh:
            confirm_response = client.post(
                "/upload",
                data={"confirm_duplicate_update": "true"},
                files={
                    "file": (
                        "sample_upload.docx",
                        fh,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=auth_headers,
            )
        assert confirm_response.status_code == 200
        confirm_payload = confirm_response.json()
        assert confirm_payload["duplicate"] is True
        assert confirm_payload["candidate_id"] == candidate_id
        assert confirm_payload["requires_confirmation"] is False
        assert confirm_payload["updated"] is True

        with TestingSessionLocal() as session:
            assert session.query(Candidate).count() == 1

            doc = session.query(Document).filter(Document.candidate_id == candidate_id).first()
            assert doc is not None
            session.add(
                Notification(
                    candidate_id=candidate_id,
                    document_id=doc.document_id,
                    message="Документ просрочен: Passport.",
                    sent=False,
                )
            )
            session.commit()

        with sample_file.open("rb") as fh:
            merge_response = client.post(
                "/upload",
                data={"confirm_duplicate_update": "true"},
                files={
                    "file": (
                        "sample_upload.docx",
                        fh,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=auth_headers,
            )
        assert merge_response.status_code == 200, merge_response.text
        assert merge_response.json()["updated"] is True

        with TestingSessionLocal() as session:
            stale = (
                session.query(Notification)
                .filter(Notification.candidate_id == candidate_id, Notification.document_id.isnot(None))
                .count()
            )
            assert stale == 0
    finally:
        app.dependency_overrides.clear()
