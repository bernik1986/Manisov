from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

import pytest
from docx import Document as DocxDocument
from sqlalchemy import Date

from models.db import SessionLocal, init_db
from models.schema import Application, Candidate, Certificate, Document, SeaService
from parser.docx_parser import DocxParser


def _make_sample_docx(path: Path) -> Path:
    doc = DocxDocument()

    personal = doc.add_table(rows=4, cols=2)
    personal.cell(0, 0).text = "Surname"
    personal.cell(0, 1).text = "Ivanov"
    personal.cell(1, 0).text = "First Name"
    personal.cell(1, 1).text = "Petr"
    personal.cell(2, 0).text = "Date of Birth"
    personal.cell(2, 1).text = "1990-01-01"
    personal.cell(3, 0).text = "EMAIL"
    personal.cell(3, 1).text = "p.ivanov@example.com"

    documents = doc.add_table(rows=2, cols=4)
    documents.cell(0, 0).text = "Passport"
    documents.cell(0, 1).text = "Passport No."
    documents.cell(0, 2).text = "Date of Issue"
    documents.cell(0, 3).text = "Date of Expiry"
    documents.cell(1, 0).text = "Passport"
    documents.cell(1, 1).text = "AA1234567"
    documents.cell(1, 2).text = "2020-01-01"
    documents.cell(1, 3).text = "2030-01-01"

    certificates = doc.add_table(rows=2, cols=3)
    certificates.cell(0, 0).text = "Certificate"
    certificates.cell(0, 1).text = "Certificate No"
    certificates.cell(0, 2).text = "Expiry Date"
    certificates.cell(1, 0).text = "COC"
    certificates.cell(1, 1).text = "COC-7788"
    certificates.cell(1, 2).text = "2029-12-31"

    sea_service = doc.add_table(rows=2, cols=3)
    sea_service.cell(0, 0).text = "Vessel name"
    sea_service.cell(0, 1).text = "Rank"
    sea_service.cell(0, 2).text = "Sign on"
    sea_service.cell(1, 0).text = "SEA STAR"
    sea_service.cell(1, 1).text = "2/O"
    sea_service.cell(1, 2).text = "2023-06-01"

    application = doc.add_table(rows=2, cols=3)
    application.cell(0, 0).text = "Position to Apply for"
    application.cell(0, 1).text = "Date Applied"
    application.cell(0, 2).text = "Last Salary (USD)"
    application.cell(1, 0).text = "Second Officer"
    application.cell(1, 1).text = "2024-01-15"
    application.cell(1, 2).text = "4500"

    doc.save(path)
    return path


def _columns(model: Any) -> set[str]:
    return {column.name for column in model.__table__.columns}


def _filter_model_data(data: dict[str, Any], model: Any) -> dict[str, Any]:
    valid = _columns(model)
    return {key: value for key, value in data.items() if key in valid}


def _coerce_model_types(data: dict[str, Any], model: Any) -> dict[str, Any]:
    coerced: dict[str, Any] = {}
    model_columns = {column.name: column for column in model.__table__.columns}

    for key, value in data.items():
        column = model_columns.get(key)
        if column is None or value in (None, ""):
            coerced[key] = value
            continue

        if isinstance(column.type, Date) and isinstance(value, str):
            try:
                coerced[key] = date.fromisoformat(value)
            except ValueError:
                coerced[key] = value
        else:
            coerced[key] = value
    return coerced


def _persist_parsed_payload(parsed: dict[str, Any]) -> int:
    init_db()
    session = SessionLocal()
    try:
        candidate_data = _coerce_model_types(_filter_model_data(parsed["personal_data"], Candidate), Candidate)
        candidate = Candidate(**candidate_data)
        session.add(candidate)
        session.flush()

        for app in parsed["applications"]:
            app_data = _coerce_model_types(_filter_model_data(app, Application), Application)
            app_data["candidate_id"] = candidate.candidate_id
            session.add(Application(**app_data))

        for doc in parsed["documents"]:
            doc_data = _coerce_model_types(_filter_model_data(doc, Document), Document)
            doc_data["candidate_id"] = candidate.candidate_id
            if "document_type" not in doc_data:
                doc_data["document_type"] = doc.get("document_type") or "Unknown document"
            session.add(Document(**doc_data))

        for cert in parsed["certificates"]:
            cert_data = _coerce_model_types(_filter_model_data(cert, Certificate), Certificate)
            cert_data["candidate_id"] = candidate.candidate_id
            if "certificate_type" not in cert_data:
                cert_data["certificate_type"] = cert.get("certificate_type") or "Unknown certificate"
            session.add(Certificate(**cert_data))

        for service in parsed["sea_service"]:
            service_data = _coerce_model_types(_filter_model_data(service, SeaService), SeaService)
            service_data["candidate_id"] = candidate.candidate_id
            session.add(SeaService(**service_data))

        session.commit()
        return candidate.candidate_id
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()


@pytest.fixture
def sample_docx_file(tmp_path: Path) -> Path:
    return _make_sample_docx(tmp_path / "sample_application.docx")


def test_docx_parser_extracts_key_sections(sample_docx_file: Path) -> None:
    parser = DocxParser()
    parsed = parser.parse(sample_docx_file)

    assert parsed["personal_data"]["surname"] == "Ivanov"
    assert parsed["personal_data"]["first_name"] == "Petr"
    assert parsed["personal_data"]["email"] == "p.ivanov@example.com"
    assert isinstance(parsed["documents"], list)
    assert isinstance(parsed["certificates"], list)
    assert isinstance(parsed["sea_service"], list)
    assert isinstance(parsed["applications"], list)


def test_docx_parsed_data_saves_to_db_without_errors(sample_docx_file: Path) -> None:
    parser = DocxParser()
    parsed = parser.parse(sample_docx_file)

    candidate_id = _persist_parsed_payload(parsed)
    assert isinstance(candidate_id, int)
    assert candidate_id > 0
