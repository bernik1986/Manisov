"""Tests for scan attachment display filenames."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app import main as main_module
from app.attachment_naming import attachment_download_filename, build_scan_filename
from app.main import app, get_db_session
from models.db import Base
from models.schema import Attachment, Candidate, Certificate, Document, Role, TemplateFile, TemplateFolder, User
from passlib.context import CryptContext

pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")


@pytest.fixture
def db_setup():
    engine = __import__("sqlalchemy").create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    Base.metadata.create_all(bind=engine)
    session = TestingSessionLocal()
    try:
        yield lambda: TestingSessionLocal()
    finally:
        session.close()
        Base.metadata.drop_all(bind=engine)


def test_build_scan_filename_rank_surname_slot():
    candidate = Candidate(
        current_rank="Chief Officer",
        surname="Chernov",
        first_name="Petr",
    )
    name = build_scan_filename(candidate, "AFF", suffix=".pdf")
    assert name == "CO Chernov AFF.pdf"


def test_resolve_type_from_document_relation_uses_document_code(db_setup):
    db_session = db_setup()
    candidate = Candidate(surname="Chernov", first_name="User", current_rank="Master")
    db_session.add(candidate)
    db_session.flush()
    doc = Document(
        candidate_id=candidate.candidate_id,
        document_type="Travel Passport",
        document_category="TP",
    )
    db_session.add(doc)
    db_session.flush()
    attachment = Attachment(
        candidate_id=candidate.candidate_id,
        file_name="old.pdf",
        file_path="/tmp/old.pdf",
        description=f"document:{doc.document_id}",
    )
    name = attachment_download_filename(db_session, candidate, attachment)
    assert name == "MST Chernov TP.pdf"
    db_session.close()


def test_resolve_certificate_uses_display_code(db_setup):
    db_session = db_setup()
    candidate = Candidate(surname="Chernov", current_rank="Chief Officer")
    db_session.add(candidate)
    db_session.flush()
    cert = Certificate(
        candidate_id=candidate.candidate_id,
        certificate_type="Advanced fire fighting",
        certificate_code="AFF",
        certificate_name_raw="AFF",
    )
    db_session.add(cert)
    db_session.flush()
    attachment = Attachment(
        candidate_id=candidate.candidate_id,
        file_name="old.pdf",
        file_path="/tmp/old.pdf",
        description=f"certificate:{cert.certificate_id}",
    )
    name = attachment_download_filename(db_session, candidate, attachment)
    assert name == "CO Chernov AFF.pdf"
    db_session.close()


@pytest.fixture
def client(db_setup, tmp_path):
    templates_dir = tmp_path / "templates"
    generated_dir = tmp_path / "generated"
    manager_dir = tmp_path / "manager_files"
    podacha_dir = templates_dir / "Podacha"
    podacha_dir.mkdir(parents=True)
    generated_dir.mkdir(parents=True)
    manager_dir.mkdir(parents=True)

    template_path = podacha_dir / "info_list_new.docx"
    doc = DocxDocument()
    doc.add_paragraph("{{ rank }}")
    doc.save(template_path)

    main_module.TEMPLATES_DIR = templates_dir
    main_module.GENERATED_DIR = generated_dir
    main_module.TEMPLATES_MANAGER_DIR = manager_dir
    main_module.UPLOADS_DIR = tmp_path / "uploads"
    main_module.UPLOADS_DIR.mkdir(parents=True)

    def override_get_db_session():
        session = db_setup()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides[get_db_session] = override_get_db_session
    try:
        with TestClient(app) as test_client:
            db_session = db_setup()
            try:
                main_module.ensure_podacha_builtin_templates(
                    db_session,
                    templates_dir=templates_dir,
                    templates_manager_dir=manager_dir,
                    get_or_create_root=main_module._get_or_create_templates_root,
                )
            finally:
                db_session.close()
            yield test_client
    finally:
        app.dependency_overrides.clear()


def test_upload_attachment_sets_display_filename(client: TestClient, db_setup, tmp_path, monkeypatch):
    monkeypatch.setattr(main_module, "UPLOADS_DIR", tmp_path / "uploads")
    main_module.UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

    db_session = db_setup()
    admin_role = Role(name="admin")
    db_session.add(admin_role)
    db_session.flush()
    db_session.add(
        User(
            username="admin_naming",
            password_hash=pwd_context.hash("admin123"),
            role_id=admin_role.role_id,
            is_active=True,
        )
    )
    candidate = Candidate(surname="Chernov", first_name="Ivan", current_rank="Captain")
    db_session.add(candidate)
    db_session.flush()
    doc = Document(
        candidate_id=candidate.candidate_id,
        document_type="Seaman's Book",
        document_category="SB",
    )
    db_session.add(doc)
    db_session.commit()
    cid = candidate.candidate_id
    did = doc.document_id
    db_session.close()

    login = client.post("/auth/login", json={"username": "admin_naming", "password": "admin123"})
    headers = {"Authorization": f"Bearer {login.json()['access_token']}"}
    response = client.post(
        f"/candidates/{cid}/attachments",
        headers=headers,
        files={"file": ("scan.pdf", b"%PDF-1.4", "application/pdf")},
        data={"attachment_type": "document", "relation_id": str(did)},
    )
    assert response.status_code == 200
    assert response.json()["attachment"]["file_name"] == "MST Chernov SB.pdf"


def test_submission_zip_uses_scan_display_name(client: TestClient, db_setup, tmp_path, monkeypatch):
    monkeypatch.setattr(main_module, "UPLOADS_DIR", tmp_path / "uploads")
    main_module.UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

    db_session = db_setup()
    admin_role = Role(name="admin")
    viewer_role = Role(name="viewer")
    db_session.add_all([admin_role, viewer_role])
    db_session.flush()
    db_session.add(
        User(
            username="viewer_zip",
            password_hash=pwd_context.hash("viewer123"),
            role_id=viewer_role.role_id,
            is_active=True,
        )
    )
    candidate = Candidate(surname="Chernov", first_name="Oleg", current_rank="Second Officer")
    db_session.add(candidate)
    db_session.flush()
    cert = Certificate(
        candidate_id=candidate.candidate_id,
        certificate_type="ECDIS",
        certificate_code="ECDIS",
        certificate_name_raw="ECDIS",
    )
    db_session.add(cert)
    db_session.flush()

    scan_path = main_module.UPLOADS_DIR / f"{uuid4().hex}.pdf"
    scan_path.write_bytes(b"%PDF-1.4 test")
    attachment = Attachment(
        candidate_id=candidate.candidate_id,
        file_name="random.pdf",
        file_path=str(scan_path),
        description=f"certificate:{cert.certificate_id}",
    )
    db_session.add(attachment)
    db_session.flush()

    stored = f"{uuid4().hex}.docx"
    target = main_module.TEMPLATES_MANAGER_DIR / stored
    target.write_bytes((main_module.TEMPLATES_DIR / "Podacha" / "info_list_new.docx").read_bytes())
    root = TemplateFolder(name="Templates", parent_id=None)
    db_session.add(root)
    db_session.flush()
    folder = TemplateFolder(name="Podacha", parent_id=root.folder_id)
    db_session.add(folder)
    db_session.flush()
    tpl = TemplateFile(
        folder_id=folder.folder_id,
        file_name="info_list_new.docx",
        file_type="docx",
        stored_name=stored,
        relative_path=stored,
        file_size_bytes=target.stat().st_size,
    )
    db_session.add(tpl)
    db_session.commit()

    cid = candidate.candidate_id
    att_id = attachment.attachment_id
    tpl_id = tpl.template_file_id
    db_session.close()

    login = client.post("/auth/login", json={"username": "viewer_zip", "password": "viewer123"})
    headers = {"Authorization": f"Bearer {login.json()['access_token']}"}
    response = client.post(
        f"/candidates/{cid}/submission-pack",
        json={"template_file_ids": [tpl_id], "attachment_ids": [att_id]},
        headers=headers,
    )
    assert response.status_code == 200
    with zipfile.ZipFile(io.BytesIO(response.content)) as zf:
        names = zf.namelist()
        assert "2O Chernov ECDIS.pdf" in names
