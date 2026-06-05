from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

import app.main as main_module
from app.main import app, get_db_session, pwd_context
from models.db import Base
from models.schema import Application, Candidate, Certificate, Document, Role, TemplateFile, TemplateFolder, User
from scripts.save_to_db import save_candidate


@pytest.fixture
def db_setup():
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    testing_session_local = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    Base.metadata.create_all(bind=engine)
    return testing_session_local


@pytest.fixture
def db_session(db_setup):
    session = db_setup()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture
def users_fixture(db_session):
    admin_role = Role(name="admin", description="Admin")
    recruiter_role = Role(name="recruiter", description="Recruiter")
    viewer_role = Role(name="viewer", description="Viewer")
    db_session.add_all([admin_role, recruiter_role, viewer_role])
    db_session.flush()

    users = [
        User(
            username="admin_test",
            password_hash=pwd_context.hash("admin123"),
            full_name="Admin User",
            role_id=admin_role.role_id,
            is_active=True,
        ),
        User(
            username="recruiter_test",
            password_hash=pwd_context.hash("recruit123"),
            full_name="Recruiter User",
            role_id=recruiter_role.role_id,
            is_active=True,
        ),
        User(
            username="viewer_test",
            password_hash=pwd_context.hash("viewer123"),
            full_name="Viewer User",
            role_id=viewer_role.role_id,
            is_active=True,
        ),
    ]
    db_session.add_all(users)
    db_session.commit()
    return {"admin": users[0], "recruiter": users[1], "viewer": users[2]}


@pytest.fixture
def client(db_setup, tmp_path):
    templates_dir = tmp_path / "templates"
    generated_dir = tmp_path / "generated"
    templates_dir.mkdir(parents=True, exist_ok=True)
    generated_dir.mkdir(parents=True, exist_ok=True)
    manager_dir = tmp_path / "manager_files"
    manager_dir.mkdir(parents=True, exist_ok=True)
    main_module.TEMPLATES_MANAGER_DIR = manager_dir

    template_path = templates_dir / "candidate_test.docx"
    doc = DocxDocument()
    doc.add_paragraph("Candidate: {{ surname }} {{ first_name }}")
    doc.add_paragraph("Certificate: {{ certificates[0].certificate_name_raw if certificates else 'N/A' }}")
    doc.save(template_path)

    main_module.TEMPLATES_DIR = templates_dir
    main_module.GENERATED_DIR = generated_dir

    def override_get_db_session():
        session = db_setup()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides[get_db_session] = override_get_db_session
    try:
        yield TestClient(app)
    finally:
        app.dependency_overrides.clear()


def _login_header(client: TestClient, username: str, password: str) -> dict[str, str]:
    response = client.post("/auth/login", json={"username": username, "password": password})
    assert response.status_code == 200
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def test_save_candidate_creates_candidate_and_related_rows(db_session):
    payload = {
        "personal_data": {"surname": "Ivanov", "first_name": "Petr", "date_of_birth": "1990-01-01"},
        "documents": [{"document_type": "Passport", "document_number": "AA123456", "date_of_expiry": "2030-01-01"}],
        "certificates": [{"certificate_type": "COC", "certificate_name_raw": "COC", "expiry_date": "2029-12-31"}],
        "sea_service": [{"vessel_name": "SEA STAR", "rank_on_vessel": "2/O", "sign_on_date": "2023-06-01"}],
        "family_contacts": [{"full_name": "Anna Ivanova", "relationship_to_candidate": "Spouse", "phone": "+12345"}],
    }

    candidate = save_candidate(payload, db_session)
    assert candidate.candidate_id is not None
    stored = db_session.get(Candidate, candidate.candidate_id)
    assert stored is not None
    assert stored.surname == "Ivanov"
    assert len(stored.documents) == 1
    assert len(stored.certificates) == 1
    assert len(stored.sea_service) == 1
    assert len(stored.family_contacts) == 1


def test_search_and_filtering(client: TestClient, db_session, users_fixture):
    near_expiry = date.today() + timedelta(days=30)
    far_expiry = date.today() + timedelta(days=500)

    c1 = Candidate(surname="Petrov", first_name="Ivan", current_rank="Chief Officer")
    c2 = Candidate(surname="Sidorov", first_name="Nikolay", current_rank="Captain")
    db_session.add_all([c1, c2])
    db_session.flush()

    db_session.add(
        Application(
            candidate_id=c2.candidate_id,
            position_applied_for="Captain",
            rank_applied_for="Master",
        )
    )
    db_session.add(Document(candidate_id=c1.candidate_id, document_type="Passport", date_of_expiry=near_expiry))
    db_session.add(Certificate(candidate_id=c2.candidate_id, certificate_type="COC", expiry_date=far_expiry))
    db_session.commit()

    auth_header = _login_header(client, "viewer_test", "viewer123")
    by_name = client.get("/candidates/search", params={"surname": "pet"}, headers=auth_header)
    assert by_name.status_code == 200
    assert any(item["surname"] == "Petrov" for item in by_name.json()["items"])

    by_rank = client.get("/candidates/search", params={"rank": "captain"}, headers=auth_header)
    assert by_rank.status_code == 200
    assert any(item["surname"] == "Sidorov" for item in by_rank.json()["items"])

    by_warning = client.get("/candidates/search", params={"expiry_warning": True}, headers=auth_header)
    assert by_warning.status_code == 200
    assert any(item["surname"] == "Petrov" for item in by_warning.json()["items"])


def test_auth_with_different_roles(client: TestClient, db_session, users_fixture):
    candidate = Candidate(surname="Auth", first_name="Case")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)

    admin_header = _login_header(client, "admin_test", "admin123")
    recruiter_header = _login_header(client, "recruiter_test", "recruit123")
    viewer_header = _login_header(client, "viewer_test", "viewer123")

    add_doc_body = {"document_type": "Passport", "document_number": "ZZ111"}

    resp_admin_add = client.post(f"/candidates/{candidate.candidate_id}/documents", json=add_doc_body, headers=admin_header)
    assert resp_admin_add.status_code == 200
    doc_id = resp_admin_add.json()["document"]["document_id"]

    resp_viewer_add = client.post(f"/candidates/{candidate.candidate_id}/documents", json=add_doc_body, headers=viewer_header)
    assert resp_viewer_add.status_code == 403

    resp_recruiter_delete = client.delete(
        f"/candidates/{candidate.candidate_id}/documents/{doc_id}",
        headers=recruiter_header,
    )
    assert resp_recruiter_delete.status_code == 200

    resp_admin_add2 = client.post(
        f"/candidates/{candidate.candidate_id}/documents", json=add_doc_body, headers=admin_header
    )
    assert resp_admin_add2.status_code == 200
    doc_id_2 = resp_admin_add2.json()["document"]["document_id"]

    resp_viewer_delete = client.delete(
        f"/candidates/{candidate.candidate_id}/documents/{doc_id_2}",
        headers=viewer_header,
    )
    assert resp_viewer_delete.status_code == 403

    resp_admin_delete = client.delete(
        f"/candidates/{candidate.candidate_id}/documents/{doc_id_2}",
        headers=admin_header,
    )
    assert resp_admin_delete.status_code == 200


def test_recruiter_can_delete_subresources_not_whole_candidate(
    client: TestClient, db_session, users_fixture, tmp_path, monkeypatch
) -> None:
    """Recruiter may DELETE nested rows and attachments; only admin may DELETE the candidate."""
    monkeypatch.setattr(main_module, "UPLOADS_DIR", tmp_path / "uploads")
    main_module.UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

    candidate = Candidate(surname="Delete", first_name="Matrix")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)
    cid = candidate.candidate_id

    admin_h = _login_header(client, "admin_test", "admin123")
    rec_h = _login_header(client, "recruiter_test", "recruit123")

    cert = client.post(
        f"/candidates/{cid}/certificates",
        json={"certificate_type": "COC"},
        headers=admin_h,
    )
    assert cert.status_code == 200
    cert_id = cert.json()["certificate"]["certificate_id"]
    assert client.delete(f"/candidates/{cid}/certificates/{cert_id}", headers=rec_h).status_code == 200

    flag_row = client.post(
        f"/candidates/{cid}/flag-documents",
        json={"flag_country": "LR"},
        headers=admin_h,
    )
    assert flag_row.status_code == 200
    flag_id = flag_row.json()["flag_document"]["flag_document_id"]
    assert client.delete(f"/candidates/{cid}/flag-documents/{flag_id}", headers=rec_h).status_code == 200

    fam = client.post(
        f"/candidates/{cid}/family-contacts",
        json={"full_name": "Contact One"},
        headers=admin_h,
    )
    assert fam.status_code == 200
    fam_id = fam.json()["family_contact"]["family_contact_id"]
    assert client.delete(f"/candidates/{cid}/family-contacts/{fam_id}", headers=rec_h).status_code == 200

    sea = client.post(
        f"/candidates/{cid}/sea-service",
        json={"vessel_name": "MV Test"},
        headers=admin_h,
    )
    assert sea.status_code == 200
    sea_id = sea.json()["sea_service"]["sea_service_id"]
    assert client.delete(f"/candidates/{cid}/sea-service/{sea_id}", headers=rec_h).status_code == 200

    up = client.post(
        f"/candidates/{cid}/attachments",
        files={"file": ("x.pdf", b"%PDF-1.4 test", "application/pdf")},
        headers=admin_h,
    )
    assert up.status_code == 200
    att_id = up.json()["attachment"]["attachment_id"]
    assert client.delete(f"/attachments/{att_id}", headers=rec_h).status_code == 200

    assert client.delete(f"/candidates/{cid}", headers=rec_h).status_code == 403
    assert client.delete(f"/candidates/{cid}", headers=admin_h).status_code == 200


def test_generate_document_from_template(client: TestClient, db_session, users_fixture):
    candidate = Candidate(surname="Template", first_name="User", current_rank="Chief Officer")
    db_session.add(candidate)
    db_session.flush()
    db_session.add(
        Certificate(
            candidate_id=candidate.candidate_id,
            certificate_type="COC",
            certificate_name_raw="Certificate of Competency",
        )
    )
    db_session.commit()

    viewer_header = _login_header(client, "viewer_test", "viewer123")
    response = client.post(f"/candidates/{candidate.candidate_id}/generate/candidate_test", headers=viewer_header)
    assert response.status_code == 200
    assert (
        response.headers.get("content-type")
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(response.content) > 0


def test_generate_document_by_template_file_id_uses_exact_managed_file(client: TestClient, db_session, users_fixture):
    """When many managed templates share the same display name, query param pins the correct row."""
    stored = f"{uuid4().hex}.docx"
    target_path = main_module.TEMPLATES_MANAGER_DIR / stored
    doc = DocxDocument()
    doc.add_paragraph("Managed: {{ surname }} {{ first_name }}")
    doc.save(target_path)

    root = TemplateFolder(name="Templates", parent_id=None)
    db_session.add(root)
    db_session.flush()
    row = TemplateFile(
        folder_id=root.folder_id,
        file_name="same_name.docx",
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        stored_name=stored,
        relative_path=stored,
        file_size_bytes=target_path.stat().st_size,
    )
    db_session.add(row)
    db_session.commit()
    db_session.refresh(row)

    candidate = Candidate(surname="Pin", first_name="Id", current_rank="CO")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)

    viewer_header = _login_header(client, "viewer_test", "viewer123")
    response = client.post(
        f"/candidates/{candidate.candidate_id}/generate/same_name.docx?template_file_id={row.template_file_id}",
        headers=viewer_header,
    )
    assert response.status_code == 200
    assert len(response.content) > 0
