from __future__ import annotations

import io

import pandas as pd
import pytest
from docx import Document as DocxDocument
from docxtpl import DocxTemplate
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

import app.main as main_module
from app.main import app, get_db_session, pwd_context
from models.db import Base
from models.schema import Company, CompanyFolder, Role, User, Vessel


@pytest.fixture
def db_setup():
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    testing_session_local = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    Base.metadata.create_all(bind=engine)
    return testing_session_local


@pytest.fixture
def db_session(db_setup):
    session = db_setup()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture
def users_fixture(db_session):
    admin_role = Role(name="admin", description="Admin")
    recruiter_role = Role(name="recruiter", description="Recruiter")
    viewer_role = Role(name="viewer", description="Viewer")
    db_session.add_all([admin_role, recruiter_role, viewer_role])
    db_session.flush()
    users = [
        User(
            username="admin_co",
            password_hash=pwd_context.hash("admin123"),
            full_name="Admin",
            role_id=admin_role.role_id,
            is_active=True,
        ),
        User(
            username="viewer_co",
            password_hash=pwd_context.hash("viewer123"),
            full_name="Viewer",
            role_id=viewer_role.role_id,
            is_active=True,
        ),
    ]
    db_session.add_all(users)
    db_session.commit()
    return {"admin": users[0], "viewer": users[1]}


@pytest.fixture
def client(db_setup, tmp_path):
    templates_dir = tmp_path / "templates"
    templates_dir.mkdir(parents=True)
    generated_dir = tmp_path / "generated"
    generated_dir.mkdir(parents=True)
    manager_dir = templates_dir / "manager_files"
    manager_dir.mkdir(parents=True)

    def override_get_db():
        db = db_setup()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db_session] = override_get_db
    main_module.TEMPLATES_DIR = templates_dir
    main_module.GENERATED_DIR = generated_dir
    main_module.TEMPLATES_MANAGER_DIR = manager_dir
    with TestClient(app) as test_client:
        yield test_client
    app.dependency_overrides.clear()


def _login_header(client: TestClient, username: str, password: str) -> dict[str, str]:
    response = client.post("/auth/login", json={"username": username, "password": password})
    assert response.status_code == 200
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def test_companies_manager_crud_and_slug_collision(client: TestClient, db_session, users_fixture):
    admin_h = _login_header(client, "admin_co", "admin123")
    viewer_h = _login_header(client, "viewer_co", "viewer123")

    listing = client.get("/companies-manager", headers=viewer_h)
    assert listing.status_code == 200
    root_id = listing.json()["root_folder_id"]

    folder_resp = client.post(
        "/companies-manager/folders",
        json={"name": "Fleet Group", "parent_id": root_id},
        headers=admin_h,
    )
    assert folder_resp.status_code == 200
    folder_id = folder_resp.json()["folder"]["folder_id"]

    company_a = client.post(
        "/companies-manager/companies",
        json={"name": "Century Shipping", "folder_id": folder_id},
        headers=admin_h,
    )
    assert company_a.status_code == 200
    company_id = company_a.json()["company"]["company_id"]
    slug_a = company_a.json()["company"]["slug"]
    assert slug_a == "century_shipping"

    company_b = client.post(
        "/companies-manager/companies",
        json={"name": "Century Shipping", "folder_id": folder_id},
        headers=admin_h,
    )
    assert company_b.status_code == 200
    assert company_b.json()["company"]["slug"] == "century_shipping_2"

    vessel_a = client.post(
        "/companies-manager/vessels",
        json={
            "company_id": company_id,
            "name": "Janina",
            "imo": "1234567",
            "flag": "Panama",
            "vessel_type": "Bulk",
        },
        headers=admin_h,
    )
    assert vessel_a.status_code == 200
    vessel_payload = vessel_a.json()["vessel"]
    assert vessel_payload["slug"] == "janina"
    assert vessel_payload["placeholders"]["imo"] == "{{ company_century_shipping_janina_imo }}"
    assert "port_of_registry" in vessel_payload["placeholders"]

    vessel_ext = client.post(
        "/companies-manager/vessels",
        json={
            "company_id": company_id,
            "name": "Gaslog Star",
            "port_of_registry": "Monrovia",
            "grt": "45000",
            "call_sign": "ABCD",
            "year_built": 2015,
        },
        headers=admin_h,
    )
    assert vessel_ext.status_code == 200
    ext = vessel_ext.json()["vessel"]
    assert ext["port_of_registry"] == "Monrovia"
    assert ext["year_built"] == 2015
    assert ext["placeholders"]["grt"] == "{{ company_century_shipping_gaslog_star_grt }}"

    vessel_b = client.post(
        "/companies-manager/vessels",
        json={"company_id": company_id, "name": "Janina"},
        headers=admin_h,
    )
    assert vessel_b.status_code == 200
    assert vessel_b.json()["vessel"]["slug"] == "janina_2"

    viewer_create = client.post(
        "/companies-manager/companies",
        json={"name": "Denied", "folder_id": folder_id},
        headers=viewer_h,
    )
    assert viewer_create.status_code == 403


def _sample_vessels_xlsx_bytes() -> bytes:
    df = pd.DataFrame(
        [
            {"Company": "Excel Shipping", "IMO": "9876543", "Vessel name": "Excel Star"},
            {"Company": "", "IMO": "", "Vessel name": "Excel Moon"},
        ]
    )
    buf = io.BytesIO()
    df.to_excel(buf, index=False)
    return buf.getvalue()


def test_companies_manager_xlsx_import(client: TestClient, users_fixture):
    admin_h = _login_header(client, "admin_co", "admin123")
    viewer_h = _login_header(client, "viewer_co", "viewer123")

    denied = client.post(
        "/companies-manager/import",
        headers=viewer_h,
        files={"file": ("vessels.xlsx", _sample_vessels_xlsx_bytes(), "application/vnd.ms-excel")},
    )
    assert denied.status_code == 403

    response = client.post(
        "/companies-manager/import",
        headers=admin_h,
        files={"file": ("vessels.xlsx", _sample_vessels_xlsx_bytes(), "application/vnd.ms-excel")},
    )
    assert response.status_code == 200, response.text
    stats = response.json()["stats"]
    assert stats["companies_created"] == 1
    assert stats["vessels_created"] == 2

    listing = client.get("/companies-manager", headers=admin_h)
    payload = listing.json()
    assert any(item["name"] == "Excel Shipping" for item in payload["companies"])
    assert any(item["name"] == "Excel Star" for item in payload["vessels"])
    assert any(item["name"] == "Excel Moon" for item in payload["vessels"])

    repeat = client.post(
        "/companies-manager/import",
        headers=admin_h,
        files={"file": ("vessels.xlsx", _sample_vessels_xlsx_bytes(), "application/vnd.ms-excel")},
    )
    assert repeat.status_code == 200
    repeat_stats = repeat.json()["stats"]
    assert repeat_stats["companies_existing"] == 1
    assert repeat_stats["vessels_skipped"] == 2


def test_company_placeholders_in_candidate_context(client: TestClient, db_session, users_fixture, tmp_path):
    from models.schema import Candidate

    admin_h = _login_header(client, "admin_co", "admin123")

    root = CompanyFolder(name="Companies", parent_id=None)
    db_session.add(root)
    db_session.flush()
    company = Company(folder_id=root.folder_id, name="Century", slug="century")
    db_session.add(company)
    db_session.flush()
    vessel = Vessel(
        company_id=company.company_id,
        name="Janina",
        slug="janina",
        imo="7654321",
        flag="Liberia",
        vessel_type="Tanker",
    )
    db_session.add(vessel)
    candidate = Candidate(surname="Test", first_name="User", current_rank="CO")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)

    context = main_module._serialize_candidate_context(candidate, db_session=db_session)
    assert context["company_century_janina_name"] == "Janina"
    assert context["company_century_janina_imo"] == "7654321"
    assert context["company_century_janina_flag"] == "Liberia"
    assert context["company_century_janina_type"] == "Tanker"
    vessel.port_of_registry = "Singapore"
    vessel.official_number = "12345"
    db_session.commit()
    context = main_module._serialize_candidate_context(candidate, db_session=db_session)
    assert context["company_century_janina_port_of_registry"] == "Singapore"
    assert context["company_century_janina_official_number"] == "12345"

    template_path = tmp_path / "vessel_placeholders.docx"
    doc = DocxDocument()
    doc.add_paragraph("{{ company_century_janina_imo }}")
    doc.save(template_path)

    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    out_path = tmp_path / "out.docx"
    tpl.save(out_path)
    rendered_doc = DocxDocument(str(out_path))
    assert "7654321" in rendered_doc.paragraphs[0].text
