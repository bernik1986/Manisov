"""Tests for candidate submission pack (ПОДАЧА) ZIP endpoint."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from openpyxl import Workbook, load_workbook
from PIL import Image
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app import main as main_module
from app import submission_pack as submission_pack_module
from app.main import app, get_db_session
from models.db import Base
from models.schema import Attachment, Candidate, Role, TemplateFile, TemplateFolder, User
from passlib.context import CryptContext

pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")


@pytest.fixture
def db_setup():
    engine = __import__("sqlalchemy").create_engine(
        "sqlite://",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    Base.metadata.create_all(bind=engine)
    session = TestingSessionLocal()
    try:
        yield lambda: TestingSessionLocal()
    finally:
        session.close()
        Base.metadata.drop_all(bind=engine)


@pytest.fixture
def users_fixture(db_setup):
    db_session = db_setup()
    admin_role = Role(name="admin")
    viewer_role = Role(name="viewer")
    db_session.add_all([admin_role, viewer_role])
    db_session.flush()
    db_session.add(
        User(
            username="admin_sub",
            password_hash=pwd_context.hash("admin123"),
            role_id=admin_role.role_id,
            is_active=True,
        )
    )
    db_session.add(
        User(
            username="viewer_sub",
            password_hash=pwd_context.hash("viewer123"),
            role_id=viewer_role.role_id,
            is_active=True,
        )
    )
    db_session.commit()
    yield
    db_session.close()


@pytest.fixture
def client(db_setup, tmp_path, users_fixture):
    templates_dir = tmp_path / "templates"
    generated_dir = tmp_path / "generated"
    manager_dir = tmp_path / "manager_files"
    podacha_dir = templates_dir / "Podacha"
    podacha_dir.mkdir(parents=True)
    generated_dir.mkdir(parents=True)
    manager_dir.mkdir(parents=True)

    template_path = podacha_dir / "info_list_new.docx"
    doc = DocxDocument()
    doc.add_paragraph("Rank {{ rank }} {{ surname }} {{ first_name }}")
    doc.add_paragraph("Opening m/v {{ opening_vessel }}")
    doc.save(template_path)

    main_module.TEMPLATES_DIR = templates_dir
    main_module.GENERATED_DIR = generated_dir
    main_module.TEMPLATES_MANAGER_DIR = manager_dir
    main_module.UPLOADS_DIR = tmp_path / "uploads"
    main_module.UPLOADS_DIR.mkdir(parents=True)

    def override_get_db_session():
        session = db_setup()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides[get_db_session] = override_get_db_session
    try:
        with TestClient(app) as test_client:
            db_session = db_setup()
            try:
                main_module.ensure_podacha_builtin_templates(
                    db_session,
                    templates_dir=templates_dir,
                    templates_manager_dir=manager_dir,
                    get_or_create_root=main_module._get_or_create_templates_root,
                )
            finally:
                db_session.close()
            yield test_client
    finally:
        app.dependency_overrides.clear()


def _login(client: TestClient, username: str, password: str) -> dict[str, str]:
    response = client.post("/auth/login", json={"username": username, "password": password})
    assert response.status_code == 200
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def _photo_bytes(color: str = "navy") -> bytes:
    output = io.BytesIO()
    Image.new("RGB", (240, 320), color).save(output, format="PNG")
    return output.getvalue()


def _large_scan_pdf() -> bytes:
    image = Image.effect_noise((1400, 1800), 90).convert("RGB")
    output = io.BytesIO()
    image.save(output, format="PDF", quality=95, resolution=150)
    image.close()
    return output.getvalue()


def test_candidate_photo_upload_preview_replace_and_delete(client: TestClient, db_setup):
    db_session = db_setup()
    candidate = Candidate(surname="Portrait", first_name="User")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)
    candidate_id = candidate.candidate_id
    headers = _login(client, "admin_sub", "admin123")

    first = client.post(
        f"/candidates/{candidate_id}/photo",
        files={"file": ("portrait.png", _photo_bytes(), "image/png")},
        headers=headers,
    )
    assert first.status_code == 200
    first_id = first.json()["photo"]["attachment_id"]
    stored = db_session.get(Attachment, first_id)
    first_path = Path(stored.file_path)
    assert first_path.is_file()
    assert stored.file_type == "image/jpeg"

    detail = client.get(f"/candidates/{candidate_id}", headers=headers)
    assert detail.status_code == 200
    assert detail.json()["photo"]["attachment_id"] == first_id
    assert all(item.get("source") != "candidate_photo" for item in detail.json()["attachments"])

    preview = client.get(f"/candidates/{candidate_id}/photo", headers=headers)
    assert preview.status_code == 200
    assert preview.headers["content-type"] == "image/jpeg"
    assert preview.content.startswith(b"\xff\xd8")

    second = client.post(
        f"/candidates/{candidate_id}/photo",
        files={"file": ("replacement.jpg", _photo_bytes("green"), "image/jpeg")},
        headers=headers,
    )
    assert second.status_code == 200
    db_session.expire_all()
    assert db_session.query(Attachment).filter(Attachment.source == "candidate_photo").count() == 1
    assert not first_path.exists()

    deleted = client.delete(f"/candidates/{candidate_id}/photo", headers=headers)
    assert deleted.status_code == 200
    assert client.get(f"/candidates/{candidate_id}/photo", headers=headers).status_code == 404
    db_session.close()


def test_submission_pack_can_contain_candidate_photo_only(client: TestClient, db_setup):
    db_session = db_setup()
    candidate = Candidate(surname="Portrait", first_name="Pack")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)
    candidate_id = candidate.candidate_id
    headers = _login(client, "admin_sub", "admin123")

    uploaded = client.post(
        f"/candidates/{candidate_id}/photo",
        files={"file": ("portrait.png", _photo_bytes(), "image/png")},
        headers=headers,
    )
    assert uploaded.status_code == 200
    response = client.post(
        f"/candidates/{candidate_id}/submission-pack",
        json={"template_file_ids": [], "attachment_ids": [], "include_candidate_photo": True},
        headers=headers,
    )
    assert response.status_code == 200
    with zipfile.ZipFile(io.BytesIO(response.content)) as archive:
        names = archive.namelist()
        assert names == ["PHOTO_PORTRAIT_PACK.jpg"]
        assert archive.read(names[0]).startswith(b"\xff\xd8")
    db_session.close()


def test_submission_pack_endpoint_enforces_archive_limit(client: TestClient, db_setup, monkeypatch):
    monkeypatch.setattr(submission_pack_module, "MAX_SUBMISSION_ZIP_BYTES", 450_000)
    db_session = db_setup()
    candidate = Candidate(surname="Limited", first_name="Archive", current_rank="Master")
    db_session.add(candidate)
    db_session.flush()
    scan_bytes = _large_scan_pdf()
    attachments = []
    for index in range(2):
        path = main_module.UPLOADS_DIR / f"large-scan-{index}.pdf"
        path.write_bytes(scan_bytes)
        attachment = Attachment(
            candidate_id=candidate.candidate_id,
            file_name=f"large-scan-{index}.pdf",
            file_type="application/pdf",
            file_path=str(path),
            file_size_bytes=len(scan_bytes),
            source=f"scan_{index}",
        )
        db_session.add(attachment)
        attachments.append(attachment)
    db_session.commit()
    candidate_id = candidate.candidate_id
    attachment_ids = [item.attachment_id for item in attachments]

    headers = _login(client, "admin_sub", "admin123")
    response = client.post(
        f"/candidates/{candidate_id}/submission-pack",
        json={"template_file_ids": [], "attachment_ids": attachment_ids},
        headers=headers,
    )

    assert response.status_code == 200
    assert len(response.content) <= 450_000
    with zipfile.ZipFile(io.BytesIO(response.content)) as archive:
        assert len(archive.namelist()) == 2
        assert all(archive.read(name).startswith(b"%PDF") for name in archive.namelist())
    db_session.close()


def test_submission_pack_returns_zip_with_generated_docx(client: TestClient, db_setup):
    db_session = db_setup()
    candidate = Candidate(
        surname="Martynyuk",
        first_name="Gennadiy",
        current_rank="Master",
        watch_officer_since_year=2015,
        home_airport="Warsaw, Poland",
        english_level="Normal",
    )
    db_session.add(candidate)
    db_session.flush()

    stored = f"{uuid4().hex}.docx"
    src = main_module.TEMPLATES_DIR / "Podacha" / "info_list_new.docx"
    target = main_module.TEMPLATES_MANAGER_DIR / stored
    target.write_bytes(src.read_bytes())

    root = TemplateFolder(name="Templates", parent_id=None)
    db_session.add(root)
    db_session.flush()
    folder = TemplateFolder(name="Podacha", parent_id=root.folder_id)
    db_session.add(folder)
    db_session.flush()
    tpl = TemplateFile(
        folder_id=folder.folder_id,
        file_name="info_list_new.docx",
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        stored_name=stored,
        relative_path=stored,
        file_size_bytes=target.stat().st_size,
    )
    db_session.add(tpl)
    db_session.commit()
    db_session.refresh(candidate)
    db_session.refresh(tpl)

    headers = _login(client, "viewer_sub", "viewer123")
    response = client.post(
        f"/candidates/{candidate.candidate_id}/submission-pack",
        json={
            "opening_vessel": "Edessaikos",
            "previous_vessel": "",
            "template_file_ids": [tpl.template_file_id],
            "attachment_ids": [],
        },
        headers=headers,
    )
    assert response.status_code == 200
    assert response.headers.get("content-type") == "application/zip"
    assert len(response.content) > 0

    with zipfile.ZipFile(io.BytesIO(response.content)) as zf:
        names = zf.namelist()
        assert len(names) == 1
        assert names[0].lower().endswith(".docx")

    db_session.close()


def test_submission_pack_can_include_generated_xlsx(client: TestClient, db_setup):
    db_session = db_setup()
    candidate = Candidate(surname="Excelpack", first_name="User", current_rank="Master")
    db_session.add(candidate)
    db_session.flush()

    stored = f"{uuid4().hex}.xlsx"
    target = main_module.TEMPLATES_MANAGER_DIR / stored
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "{{ surname }}"
    sheet["B1"] = "{{ rank }}"
    workbook.save(target)

    root = TemplateFolder(name="Templates", parent_id=None)
    db_session.add(root)
    db_session.flush()
    folder = TemplateFolder(name="Podacha", parent_id=root.folder_id)
    db_session.add(folder)
    db_session.flush()
    tpl = TemplateFile(
        folder_id=folder.folder_id,
        file_name="info_list.xlsx",
        file_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        stored_name=stored,
        relative_path=stored,
        file_size_bytes=target.stat().st_size,
    )
    db_session.add(tpl)
    db_session.commit()
    db_session.refresh(candidate)
    db_session.refresh(tpl)

    headers = _login(client, "viewer_sub", "viewer123")
    response = client.post(
        f"/candidates/{candidate.candidate_id}/submission-pack",
        json={"template_file_ids": [tpl.template_file_id], "attachment_ids": []},
        headers=headers,
    )

    assert response.status_code == 200
    with zipfile.ZipFile(io.BytesIO(response.content)) as zf:
        names = zf.namelist()
        assert len(names) == 1
        assert names[0].lower().endswith(".xlsx")
        rendered = load_workbook(io.BytesIO(zf.read(names[0])), data_only=False)
        sheet = rendered.active
        assert sheet["A1"].value == "EXCELPACK"
        assert sheet["B1"].value == "Master"

    db_session.close()


def test_submission_pack_rejects_empty_selection(client: TestClient, db_setup):
    db_session = db_setup()
    candidate = Candidate(surname="Empty", first_name="Pack")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)

    headers = _login(client, "admin_sub", "admin123")
    response = client.post(
        f"/candidates/{candidate.candidate_id}/submission-pack",
        json={"template_file_ids": [], "attachment_ids": []},
        headers=headers,
    )
    assert response.status_code == 400
    db_session.close()


def test_candidate_update_persists_submission_fields(client: TestClient, db_setup):
    db_session = db_setup()
    candidate = Candidate(surname="Fields", first_name="Test")
    db_session.add(candidate)
    db_session.commit()
    db_session.refresh(candidate)

    headers = _login(client, "admin_sub", "admin123")
    response = client.put(
        f"/candidates/{candidate.candidate_id}",
        json={
            "home_airport": "Odessa",
            "desirable_salary_usd": 9500,
            "rejoin_bonus_usd": 500,
            "vaccination_summary": "fully vaccinated",
        },
        headers=headers,
    )
    assert response.status_code == 200
    db_session.expire_all()
    stored = db_session.get(Candidate, candidate.candidate_id)
    assert stored.home_airport == "Odessa"
    assert stored.desirable_salary_usd == 9500
    assert stored.rejoin_bonus_usd == 500
    assert stored.vaccination_summary == "fully vaccinated"
    db_session.close()
